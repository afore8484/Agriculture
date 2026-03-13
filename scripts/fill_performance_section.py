from copy import deepcopy
from pathlib import Path
import json
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\workspaces\Agriculture\prototype")
SOURCE_DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.docx"
TARGET_DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.performance_updated.docx"
BACKUP_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.before_performance_fill.docx"
PAYLOAD_PATH = BASE / "tmp" / "performance_payload.json"


def insert_paragraph_after(anchor, text=""):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        para.add_run(text)
    return para


def clone_table_after(doc, anchor, data):
    sample = doc.tables[0]
    new_tbl = deepcopy(sample._tbl)
    for tr in list(new_tbl.tr_lst):
        new_tbl.remove(tr)
    row_template = deepcopy(sample.rows[0]._tr)
    for _ in data:
        new_tbl.append(deepcopy(row_template))
    anchor._p.addnext(new_tbl)
    table = Table(new_tbl, anchor._parent)
    table.style = sample.style
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row[:4]):
            table.cell(r_idx, c_idx).text = str(value)
    return table


def delete_between(p_start, p_end):
    cur = p_start._p.getnext()
    while cur is not None and cur != p_end._p:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def main():
    payload = json.loads(PAYLOAD_PATH.read_text(encoding="utf-8-sig"))
    if not BACKUP_PATH.exists():
        shutil.copy2(SOURCE_DOCX_PATH, BACKUP_PATH)
    doc = Document(SOURCE_DOCX_PATH)
    paragraphs = doc.paragraphs
    i_51 = next(i for i,p in enumerate(paragraphs) if i > 150 and p.text.strip().startswith("5.1 "))
    i_52 = next(i for i,p in enumerate(paragraphs) if i > i_51 and p.text.strip().startswith("5.2 "))
    p51 = paragraphs[i_51]
    p52 = paragraphs[i_52]
    delete_between(p51, p52)
    p = insert_paragraph_after(p51, payload["intro"])
    clone_table_after(doc, p, payload["rows"])
    insert_paragraph_after(p, payload["note"])
    doc.save(TARGET_DOCX_PATH)
    print(TARGET_DOCX_PATH)


if __name__ == "__main__":
    main()
