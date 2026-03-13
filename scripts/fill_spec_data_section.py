#!/usr/bin/env python3
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

DOC_PATH = Path(r"E:\workspaces\Agriculture\prototype\output\doc\spec_analysis_draft_v1.docx")
BACKUP_PATH = DOC_PATH.with_name("spec_analysis_draft_v1.before_data_fill.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(doc: Document, paragraph: Paragraph, rows: list[list[str]], style: str = "Table Grid") -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    paragraph._p.addnext(table._tbl)
    return table


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def insert_items(doc: Document, anchor: Paragraph, items: list[tuple[str, object]]) -> None:
    for kind, payload in reversed(items):
        if kind == "paragraph":
            insert_paragraph_after(anchor, payload)
        elif kind == "table":
            insert_table_after(doc, anchor, payload)
        else:
            raise ValueError(f"Unsupported item kind: {kind}")


def main() -> None:
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(str(DOC_PATH))

    p4 = find_paragraph(doc, "4. 数据需求")
    p41 = find_paragraph(doc, "4.1 数据分类")
    p42 = find_paragraph(doc, "4.2 数据库设计概述")

    for paragraph in list(doc.paragraphs):
        if "【待补充】" in paragraph.text:
            remove_paragraph(paragraph)

    insert_paragraph_after(
        p4,
        "本章依据 docs/data 目录下各模块数据结构设计文档整理，补充平台级数据分类、存储选型和核心数据表分布。详细字段定义、主键约束、索引策略和扩展字段设计以各模块数据结构设计文档为准。",
    )

    items_41 = [
        ("paragraph", "平台数据按“静态数据”和“动态数据”两类管理。静态数据承担统一口径、组织关联、权限控制和规则配置作用；动态数据承载业务办理、状态流转、统计快照和审计留痕。"),
        ("paragraph", "4.1.1 静态数据"),
        ("paragraph", "静态数据主要来源于组织区划、权限体系、字典配置、财务基础设置和各模块主档模板，是全平台统一编码、统一口径和统一授权的基础。"),
        ("table", [
            ["数据类别", "来源模块", "典型对象/数据表", "主要用途"],
            ["行政区划与组织主数据", "总平台、农经一张图、绩效、审计、收入台账、监管", "adm_region、dim_region、sys_org、dim_org", "统一区划口径、组织归属、区域汇总和地图定位"],
            ["用户、角色与权限主数据", "村委财务、绩效、审计、监管", "sys_user、sys_role、sys_user_role、sys_permission", "登录鉴权、菜单授权、流程分派和数据范围控制"],
            ["数据字典与规则配置", "全平台", "sys_dict_type、sys_dict_item、warning_rule、指标/状态字典", "统一状态枚举、业务规则、预警规则和页面下拉选项"],
            ["财务基础主数据", "村委财务、农经一张图", "账套、会计期间、会计科目、科目类别、统计指标定义", "统一会计口径、核算维度和统计分析口径"],
            ["业务主档与模板", "财务、审计、绩效、监管", "资产分类、合同模板、审计事项库、评分指标、报告模板、流程模板", "支撑业务初始化、模板复用和流程配置"],
            ["地图与分析维度", "农经一张图", "dim_time、dim_region、dim_org", "支撑地图联动、时间趋势、专题钻取和多维分析"],
        ]),
        ("paragraph", "4.1.2 动态数据"),
        ("paragraph", "动态数据随业务办理持续产生，覆盖财务核算、资产资源、合同履约、收入上报、绩效评价、智能审计、主体监管、风险预警和日志留痕等场景。"),
        ("table", [
            ["数据类别", "来源模块", "典型对象/数据表", "主要用途"],
            ["财务核算与凭证数据", "村委财务、农经一张图", "凭证主表、凭证明细表、银行日记账、现金日记账、财务月汇总表", "支持凭证录入、账簿生成、报表统计和专题汇聚"],
            ["资产、资源与合同数据", "村委财务、农经一张图", "资产卡片、资产台账、资源台账、合同主表、合同执行表", "支持资产资源管理、合同履约、收付款跟踪和专题展示"],
            ["收入台账与上报数据", "村集体经济收入台账", "收入台账主表、收入台账明细表、上报批次表、区域汇总统计表", "支持 PC/移动填报、审核上报、进度跟踪和统计分析"],
            ["绩效评价与评分数据", "村集体绩效评价管理系统", "评价事项、采集任务、指标采集值、评分结果、评价报告", "支持指标采集、审核、评分、报告输出和区县报送"],
            ["审计过程数据", "村集体智能审计系统", "审计计划、审计方案、资料调取、审计底稿、审计报告、整改跟踪", "支持审计全过程管理和问题闭环"],
            ["主体监管与项目申报数据", "新型农业经营主体财务经营监管系统", "主体档案、成员出资、项目申报、示范监测、金融产品与贷后数据", "支持主体画像、项目审批、示范监管和金融协同"],
            ["预警快照与日志数据", "全平台", "预警事件、报表快照、操作日志、登录日志、同步日志", "支持风险预警、统计快照、运行监控和审计追溯"],
        ]),
    ]
    insert_items(doc, p41, items_41)

    items_42 = [
        ("paragraph", "平台数据库设计遵循“统一主数据、分模块建模、分析与交易分层、日志留痕完整”的原则，在满足业务交易一致性的同时兼顾统计分析、地图联动和附件归档能力。"),
        ("paragraph", "4.2.1 数据库类型与选型"),
        ("table", [
            ["类型", "建议选型", "适用范围", "说明"],
            ["业务主库", "PostgreSQL 15/16 或兼容 MySQL 8", "总平台、绩效、收入台账、监管、村委财务、审计", "关系型事务存储，支持多表关联、统计聚合、层级查询和事务一致性"],
            ["GIS 扩展", "PostGIS", "农经一张图", "支撑区划边界、地图定位、空间联动和地图专题分析"],
            ["缓存层", "Redis", "全平台", "缓存区划树、字典、热点统计、权限快照和高频查询结果"],
            ["文件存储", "MinIO / OSS / S3", "全平台", "存储凭证附件、合同附件、审计材料、报表导出文件和电子回单"],
            ["检索增强", "Elasticsearch（可选）", "绩效、监管、日志检索", "用于全文检索、跨字段搜索和日志快速定位"],
        ]),
        ("paragraph", "4.2.2 数据表结构"),
        ("paragraph", "根据现有数据结构设计，平台数据表总体分为公共基础表、主题事实表、业务过程表、统计快照表和日志审计表五类。规格说明书中保留模块级核心表清单，详细字段设计继续以各数据结构设计文档为准。"),
        ("table", [
            ["表类别", "典型来源模块", "代表数据表/表组", "作用说明"],
            ["公共基础表", "全平台", "adm_region、sys_org、sys_user、sys_role、sys_permission、sys_dict_item", "统一组织、区划、用户、角色、权限和字典基础能力"],
            ["主题事实表", "农经一张图", "fact_finance_monthly_summary、fact_asset_item、fact_resource_item、fact_contract_main、fact_warning_event", "承接各专题汇总结果、地图展示和预警分析"],
            ["业务过程表", "村委财务、收入台账、绩效、审计、监管", "凭证、台账、采集任务、审计方案、项目申报、审批实例等", "记录业务办理过程、状态流转和责任主体"],
            ["统计快照表", "农经一张图、收入台账、绩效、财务报表", "月汇总、报表快照、区域统计、评分结果、做账进度快照", "支撑大屏展示、趋势分析、报表留痕和历史比对"],
            ["日志审计表", "全平台", "sys_login_log、sys_op_log、同步日志、操作记录、审批历史", "满足系统运维、审计追踪和问题定位需求"],
        ]),
        ("table", [
            ["序号", "模块", "核心数据表/表组", "主要承载内容"],
            ["1", "平台公共主数据", "行政区划、组织、用户、角色、权限、字典、流程模板", "作为全部业务系统统一主数据底座和授权基础"],
            ["2", "农经一张图", "dim_region、dim_org、dim_time、fact_finance_monthly_summary、fact_asset_item、fact_contract_main、warning_rule、fact_warning_event", "承接多专题汇聚、地图展示、区域钻取和风险预警"],
            ["3", "村委财务事务管理系统", "账套、会计期间、会计科目、凭证主表、凭证明细表、银行日记账、资产卡片、合同主表、审批实例表", "承接村级财务、资金、资产、合同和审批闭环业务"],
            ["4", "村集体绩效评价管理系统", "评价事项、指标定义、采集任务、指标采集值、评分结果、评价报告", "承接指标采集、审核、评分、报告生成和区县报送"],
            ["5", "村集体经济收入台账系统", "收入台账主表、收入台账明细表、台账附件表、上报批次表、区域汇总统计表", "承接收入填报、审核上报、移动端填报和统计分析"],
            ["6", "村集体智能审计系统", "审计计划、审计方案、资料调取、审计底稿、审计报告、整改跟踪、附件表", "承接审计全过程管理、问题闭环和档案留痕"],
            ["7", "新型农业经营主体财务经营监管系统", "主体档案、成员出资、经营统计、项目申报、示范监测、金融产品与贷后、公告政策表", "承接主体监管、项目审批、金融服务协同和政策服务"],
        ]),
        ("paragraph", "说明：本节用于明确需求规格层面的数据承载范围和核心表对象，不替代详细数据库设计文档。后续接口设计、开发建模和测试数据准备，均应以 docs/data 目录下对应模块的数据结构设计文档为依据。"),
    ]
    insert_items(doc, p42, items_42)

    doc.save(str(DOC_PATH))
    print(DOC_PATH)
    print(BACKUP_PATH)

if __name__ == "__main__":
    main()
