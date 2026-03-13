#!/usr/bin/env python3
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

DOC_PATH = Path(r"E:\workspaces\Agriculture\prototype\output\doc\spec_analysis_draft_v1.docx")
BACKUP_PATH = DOC_PATH.with_name("spec_analysis_draft_v1.before_env_dict_fill.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(doc: Document, paragraph: Paragraph, rows: list[list[str]], style: str = "Table Grid") -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    paragraph._p.addnext(table._tbl)
    return table


def find_paragraph(doc: Document, text: str, occurrence: int = 1) -> Paragraph:
    matched = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            matched.append(paragraph)
    if len(matched) < occurrence:
        raise RuntimeError(f"Paragraph not found: {text} occurrence={occurrence}")
    return matched[occurrence - 1]


def clear_between(start_paragraph: Paragraph, end_paragraph: Paragraph) -> None:
    body = start_paragraph._p.getparent()
    removing = False
    for child in list(body):
        if child == start_paragraph._p:
            removing = True
            continue
        if child == end_paragraph._p:
            break
        if removing:
            body.remove(child)


def insert_items(doc: Document, anchor: Paragraph, items: list[tuple[str, object]]) -> None:
    for kind, payload in reversed(items):
        if kind == "paragraph":
            insert_paragraph_after(anchor, payload)
        elif kind == "table":
            insert_table_after(doc, anchor, payload)
        else:
            raise ValueError(f"Unsupported item kind: {kind}")


def update_database_table(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers[:4] != ["类型", "建议选型", "适用范围", "说明"]:
            continue
        for row in table.rows[1:]:
            key = row.cells[0].text.strip()
            if key == "业务主库":
                row.cells[1].text = "金仓数据库（KingbaseES）"
                row.cells[2].text = "总平台、绩效、收入台账、监管、村委财务、审计"
                row.cells[3].text = "承载核心业务数据、统计数据和权限数据，满足国产化部署要求"
            elif key == "GIS 扩展":
                row.cells[1].text = "金仓空间能力 / 独立 GIS 服务"
                row.cells[2].text = "农经一张图"
                row.cells[3].text = "支撑区划边界、地图定位、空间联动和地图专题分析"


def main() -> None:
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(str(DOC_PATH))

    p231 = find_paragraph(doc, "2.3.1 硬件环境")
    p232 = find_paragraph(doc, "2.3.2 软件环境")
    p24 = find_paragraph(doc, "2.4 约束与假设")
    p43 = find_paragraph(doc, "4.3 数据词典")
    p44 = find_paragraph(doc, "4.4 数据采集与清洗规则")

    clear_between(p231, p232)
    clear_between(p232, p24)
    clear_between(p43, p44)

    hardware_items = [
        ("paragraph", "平台建议采用 5 台服务器部署，其中 3 台应用服务节点、2 台数据库节点，满足业务处理、统计分析、文件存储和数据库高可用要求。"),
        ("table", [
            ["序号", "数量", "硬件配置", "建议部署角色"],
            ["1", "3 台", "内存 32GB，CPU 16 核，SSD 1TB，普通硬盘 2TB", "应用服务、接口服务、报表任务、文件处理、定时任务"],
            ["2", "2 台", "内存 64GB，CPU 16 核，SSD 2TB，普通硬盘 1TB", "金仓数据库主备节点"],
        ]),
        ("paragraph", "硬件规划说明：应用节点负责门户访问、业务处理、批量任务和附件文件处理；数据库节点负责事务处理、统计查询、备份恢复和主备切换。"),
    ]
    insert_items(doc, p231, hardware_items)

    software_items = [
        ("paragraph", "软件环境采用国产化部署路线，操作系统使用麒麟，数据库采用金仓数据库，应用运行环境使用 Java 17 及以上版本。"),
        ("table", [
            ["类别", "配置要求", "用途说明"],
            ["操作系统", "麒麟操作系统", "用于应用服务器和数据库服务器部署，满足国产化适配要求"],
            ["运行环境", "Java 17+", "作为后端应用统一运行时环境"],
            ["数据库", "金仓数据库（KingbaseES）", "承载核心业务数据、统计数据和权限数据"],
            ["缓存", "Redis", "缓存区划树、字典、热点查询和权限快照"],
            ["对象存储", "MinIO / 兼容对象存储", "存储附件、电子回单、审计材料和报表导出文件"],
        ]),
    ]
    insert_items(doc, p232, software_items)

    dict_items = [
        ("paragraph", "数据字典用于统一全平台业务编码、状态枚举、分类口径和接口传输值，确保跨模块的数据采集、业务流转、统计分析和页面展示保持一致。"),
        ("table", [
            ["字典类别", "典型字典项", "适用模块", "用途说明"],
            ["组织区划类", "区划级别、组织类型、部门类型", "全平台", "统一组织归属、区域过滤、统计汇总和权限范围"],
            ["用户权限类", "用户状态、角色类型、权限类型、岗位类型", "村委财务、绩效、审计、监管", "支撑登录鉴权、菜单授权、流程分派和账号管理"],
            ["财务核算类", "凭证字、会计科目类别、余额方向、凭证状态、审核状态", "村委财务、农经一张图", "支撑凭证处理、账簿生成、财务统计和数据汇聚"],
            ["资产合同类", "资产类别、资产状态、合同类型、合同状态、收支类型", "村委财务、农经一张图", "支撑资产管理、合同履约、收付款统计和专题分析"],
            ["绩效评价类", "事项状态、指标类型、评分等级、审核结论、报告状态", "村集体绩效评价管理系统", "支撑指标采集、评分计算、报告输出和流程控制"],
            ["收入台账类", "收入类型、填报状态、上报状态、比较方式、预警级别", "村集体经济收入台账", "支撑填报、上报、预警和统计分析"],
            ["审计业务类", "审计阶段、问题类型、整改状态、底稿状态、报告类型", "村集体智能审计系统", "支撑审计流程、问题闭环和档案归档"],
            ["主体监管类", "主体类型、行业类型、项目状态、示范状态、金融产品类型", "新型农业经营主体财务经营监管系统", "支撑主体画像、项目申报、示范监测和金融协同"],
            ["系统运行类", "同步状态、任务状态、日志类型、附件类型", "全平台", "支撑系统监控、任务调度、日志检索和文件管理"],
        ]),
        ("paragraph", "字典管理原则：同一业务含义仅保留一个主字典编码；跨模块复用优先采用平台级字典；新增字典项时需同步评估页面展示、接口传输、统计口径和历史数据兼容。"),
    ]
    insert_items(doc, p43, dict_items)

    update_database_table(doc)
    doc.save(str(DOC_PATH))
    print(DOC_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
