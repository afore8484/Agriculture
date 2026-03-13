from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from pathlib import Path

OUTPUT_DIR = Path(r"E:/workspaces/Agriculture/prototype/output/doc")
SCRIPT_PATH = Path(r"E:/workspaces/Agriculture/prototype/scripts/generate_doc_templates.py")

SPEC_TOC = [
    "1. 引言",
    "1.1 编写目的",
    "1.2 项目背景",
    "1.2.1 参与单位",
    "1.2.2 系统关系",
    "1.3 术语与定义",
    "1.4 参考资料",
    "2. 项目总体描述",
    "2.1 项目目标",
    "2.2 用户特征",
    "2.3 运行环境",
    "2.3.1 硬件环境",
    "2.3.2 软件环境",
    "2.4 约束与假设",
    "3. 功能需求",
    "3.1 系统总体功能结构",
    "3.2 功能模块划分",
    "3.2.1 XXX系统",
    "3.2.2 XXX系统",
    "3.3 详细功能说明",
    "4. 数据需求",
    "4.1 数据分类",
    "4.1.1 静态数据",
    "4.1.2 动态数据",
    "4.2 数据库设计概述",
    "4.2.1 数据库类型与选型",
    "4.2.2 数据表结构概览",
    "4.3 数据词典",
    "4.4 数据采集与清洗规则",
    "5. 非功能需求",
    "5.1 性能需求",
    "5.1.1 响应时间",
    "5.1.2 并发能力",
    "5.1.3 数据精度",
    "5.2 安全性需求",
    "5.2.1 身份认证与授权",
    "5.2.2 数据加密与隐私保护",
    "5.2.3 审计与日志",
    "5.3 可用性与可靠性",
    "5.4 可维护性与可扩展性",
    "5.5 兼容性需求",
    "5.5.1 浏览器/终端兼容",
    "5.5.2 外部系统接口兼容",
    "6. 外部接口需求",
    "6.1 用户界面（UI）要求",
    "6.2 硬件接口",
    "6.3 软件接口",
    "6.3.1 与XXX系统接口",
    "6.3.2 与XXX平台接口",
    "6.3.3 与微信/短信平台接口",
    "7. 系统部署与运行需求",
    "7.1 部署架构",
    "7.2 灾备与恢复策略",
    "7.3 系统监控与运维支持",
]

HIGH_TOC = [
    "1. 引言",
    "1.1 编写目的",
    "1.2 项目背景",
    "1.3 定义",
    "1.4 参考资料",
    "2. 任务概述",
    "2.1 目标",
    "2.2 需求概述",
    "2.2.1 政策驱动需求",
    "2.2.2 社会需求响应",
    "2.2.3 行业发展需求",
    "2.2.4 核心业务需求",
    "2.3 条件与限制",
    "2.3.1 政策合规限制",
    "2.3.2 技术环境限制",
    "2.3.3 数据交互限制",
    "2.3.4 资源与进度限制",
    "2.3.5 业务流程限制",
    "3. 总体设计",
    "3.1 总体结构和模块外部设计",
    "3.1.1 系统总体结构",
    "3.1.2 核心模块外部设计",
    "3.1.3 模块间协作关系",
    "3.2 功能分配",
    "4. 接口设计",
    "4.1 外部接口",
    "4.1.1 页面交互接口",
    "4.1.2 XXX系统交互接口",
    "4.1.3 XXX平台交互接口",
    "4.1.4 XXX交互接口",
    "4.1.5 XXX平台交互接口",
    "4.2 内部接口",
    "4.2.1 微服务间接口交互",
    "4.2.2 事件总线",
    "4.2.3 数据库交互接口",
    "4.2.4 缓存交互接口",
    "4.2.5 文件服务器交互接口",
    "5. 数据结构设计",
    "5.1 XXX系统",
    "5.2 XXX系统",
    "6. 逻辑结构设计",
    "7. 物理结构设计",
    "8. 数据结构与程序的关系",
    "9. 运行设计",
    "9.1 运行模块的组合",
    "9.2 运行控制",
    "9.3 运行时间",
    "10. 出错处理设计",
    "10.1 出错输出信息",
    "10.2 出错处理对策",
    "11. 安全保密设计",
    "11.1 总体安全框架",
    "11.2 身份与访问控制",
    "11.3 数据安全",
    "11.4 网络与边界安全",
    "11.5 应用安全",
    "11.6 密码与密钥管理",
    "11.7 审计与合规",
    "11.8 灾备与恢复",
    "11.9 保密管理",
    "12. 维护设计",
    "12.1 维护工具链",
    "12.2 维护模块",
    "12.3 维护制度",
    "12.4 人员与培训设施",
]

DETAIL_TOC = [
    "1. 引言",
    "1.1 编写目的",
    "1.2 文档范围",
    "1.3 读者对象",
    "1.4 定义、首字母缩写词和缩略语",
    "1.5 参考资料",
    "1.6 文档概述",
    "2. 总体设计概述",
    "2.1 设计背景",
    "2.2 设计目标与约束",
    "2.3 系统架构概述",
    "2.3.1 逻辑架构",
    "2.3.2 物理架构",
    "2.3.3 技术栈选型",
    "2.4 与本系统相关的其他系统",
    "3. 系统详细设计",
    "3.x [模块A名称] 详细设计",
    "3.x.1 模块概述",
    "3.x.2 功能描述",
    "3.x.3 类/组件设计",
    "3.x.3.1 类图",
    "3.x.3.2 类说明",
    "3.x.4 业务流程设计",
    "3.x.4.1 [业务流程A]",
    "3.x.5 界面设计（可选）",
    "3.x.5.1 界面原型图",
    "3.x.5.2 界面元素说明",
    "4. 数据库设计",
    "4.1 设计原则",
    "4.2 逻辑数据模型",
    "4.3 物理数据模型",
    "4.3.1 表/集合清单",
    "4.3.2 表结构设计",
    "4.3.3 索引设计",
    "4.4 存储过程/函数设计（可选）",
    "4.5 数据字典",
    "5. 接口设计",
    "5.1 内部接口设计",
    "5.1.1 [接口名称A]",
    "5.2 外部接口设计",
    "6. 非功能性需求设计",
    "6.1 性能设计",
    "6.1.1 性能指标",
    "6.1.2 设计策略",
    "6.2 安全性设计",
    "6.2.1 安全威胁与应对",
    "6.2.2 身份认证与授权",
    "6.2.3 数据加密与脱敏",
    "6.2.4 日志与审计",
    "6.3 可靠性/可用性设计",
    "6.3.1 冗余部署",
    "6.3.2 容错与降级",
    "6.3.3 数据备份与恢复",
    "6.4 可维护性设计",
    "6.5 可扩展性设计",
    "7. 运行环境设计",
    "7.1 硬件环境",
    "7.1.1 服务器配置",
    "7.1.2 网络拓扑",
    "7.2 软件环境",
    "7.2.1 操作系统",
    "7.2.2 中间件",
    "7.2.3 依赖服务",
    "8. 部署方案",
    "8.1 部署结构图",
    "8.2 部署步骤",
    "8.3 配置说明",
    "9. 测试建议",
    "9.1 单元测试建议",
    "9.2 集成测试建议",
    "9.3 性能测试建议",
]

TEMPLATES = [
    ("需求规格说明书模板", SPEC_TOC, "本文档用于描述项目需求边界、功能需求、数据需求、非功能需求及外部接口要求。"),
    ("概要设计说明书模板", HIGH_TOC, "本文档用于描述系统总体架构、模块划分、接口设计、安全设计及运行维护设计。"),
    ("详细设计说明书模板", DETAIL_TOC, "本文档用于描述模块级详细设计、数据库设计、接口设计、非功能性设计和部署方案。"),
]

PLACEHOLDER_TEXT = "【填写说明】请结合本项目实际内容补充本节，保持中文表述，并在必要时插入流程图、原型图、表结构或接口示例。"
NOTE_TEXT = "备注：以上目录依据你提供的模板要求生成，当前为可直接编写的 Word 模板骨架。"


def set_cn_font(run, size_pt=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_cn_font(run, 10)


def parse_level(title):
    num = title.split(' ', 1)[0]
    if '3.x' in num:
        return 2
    return max(0, num.count('.') )


def build_template(doc_title, toc, purpose):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number(section)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title)
    set_cn_font(r, 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('适用项目：海南数字“三农”金融数据总平台')
    set_cn_font(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('版本：V1.0')
    set_cn_font(r, 12)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('文档说明：' + purpose)
    set_cn_font(r, 12)

    p = doc.add_paragraph()
    r = p.add_run(NOTE_TEXT)
    set_cn_font(r, 10)

    doc.add_page_break()

    p = doc.add_paragraph()
    r = p.add_run('目录')
    set_cn_font(r, 16, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in toc:
        level = parse_level(item)
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.6 * level)
        r = p.add_run(item)
        set_cn_font(r, 12)

    doc.add_page_break()

    for item in toc:
        level = parse_level(item)
        p = doc.add_paragraph()
        r = p.add_run(item)
        set_cn_font(r, max(14, 18 - min(level, 4) * 2), True)
        if level <= 1:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        p2 = doc.add_paragraph()
        r2 = p2.add_run(PLACEHOLDER_TEXT)
        set_cn_font(r2, 11)

        if '业务流程' in item:
            p3 = doc.add_paragraph()
            r3 = p3.add_run('【建议内容】流程名称、触发条件、参与角色、关键节点、异常处理、输出结果。')
            set_cn_font(r3, 11)
        elif '接口' in item:
            p3 = doc.add_paragraph()
            r3 = p3.add_run('【建议内容】接口功能、调用方/提供方、协议、入参、出参、示例、错误码。')
            set_cn_font(r3, 11)
        elif '数据库' in item or '数据' in item:
            p3 = doc.add_paragraph()
            r3 = p3.add_run('【建议内容】实体说明、字段定义、约束、索引、来源、更新规则、清洗规则。')
            set_cn_font(r3, 11)
        elif '安全' in item:
            p3 = doc.add_paragraph()
            r3 = p3.add_run('【建议内容】身份认证、授权控制、脱敏、加密、审计留痕、边界防护。')
            set_cn_font(r3, 11)

    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for name, toc, purpose in TEMPLATES:
        doc = build_template(name, toc, purpose)
        out = OUTPUT_DIR / f'{name}.docx'
        doc.save(out)
        print(out)


if __name__ == '__main__':
    main()
