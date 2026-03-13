from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r'E:/workspaces/Agriculture/prototype')
DOC_OUT = BASE / 'output' / 'doc' / 'spec_analysis_draft_v3_interfaces_batch_a.docx'
MD_NONGJING = BASE / 'docs' / 'interface_detail_nongjing.md'
MD_VILLAGE = BASE / 'docs' / 'interface_detail_village_finance.md'
CATALOG_JSON = BASE / 'data' / 'interfaces' / 'interface_catalog_seed.json'
PLACEHOLDER = '【待补充】本章节暂按模板保留，后续继续补充。'

modules = json.loads(CATALOG_JSON.read_text(encoding='utf-8'))['modules']

nongjing_interfaces = [
    ('获取行政区划树', 'GET', '/api/nongjing/common/regions', '获取区划树和层级信息，供全局筛选使用', '入参：parentCode、level；出参：regionCode、regionName、level、hasChildren', 'dim_region'),
    ('获取首页天气', 'GET', '/api/nongjing/overview/weather', '获取当前选中区划当天天气与温度', '入参：regionCode、date；出参：weatherType、temperatureMin、temperatureMax', 'fact_weather_daily'),
    ('获取首页总览卡片', 'GET', '/api/nongjing/overview/cards', '获取区县数、乡镇数、村集体数、未结账数、未建账套数等卡片', '入参：regionCode、yearMonth；出参：countyCount、townCount、orgCount、unclosedCount、noLedgerCount', 'dim_region、dim_org、fact_village_finance_monthly'),
    ('获取首页地图统计', 'GET', '/api/nongjing/overview/map-stats', '按区划返回地图着色和聚合统计结果', '入参：regionCode、yearMonth、metricCode；出参：regionName、metricValue、rankNo', 'mv_region_topic_rank_monthly、dim_region'),
    ('获取财务收入分档统计', 'GET', '/api/nongjing/finance/income-bands/stats', '返回5万/10万/20万分档数量和占比', '入参：regionCode、yearMonth、bandCode；出参：highCount、lowCount、highRate、lowRate', 'mv_finance_income_band_monthly、dim_income_band'),
    ('获取财务收入分档明细', 'GET', '/api/nongjing/finance/income-bands/list', '获取高于/低于阈值的村集体明细列表', '入参：regionCode、yearMonth、bandCode、compareType、pageNum、pageSize；出参：orgName、totalIncome、operateIncome、subsidyIncome、investIncome、otherIncome', 'mv_finance_income_band_monthly、dim_org'),
    ('导出财务收入分档明细', 'POST', '/api/nongjing/finance/income-bands/export', '异步导出收入分档明细', '入参：筛选快照、exportType；出参：taskId、status', 'mv_finance_income_band_monthly'),
    ('获取财务指标总览', 'GET', '/api/nongjing/finance/indicators/summary', '获取货币资金、应收账款、存货、固定资产、本年收益等指标', '入参：regionCode、yearMonth；出参：metricCode、metricName、value、unit', 'fact_village_finance_monthly、dim_metric'),
    ('获取资产负债变动分析', 'GET', '/api/nongjing/finance/balance-sheet/analysis', '获取资产合计、流动资产、非流动资产等变动分析', '入参：regionCode、startMonth、endMonth；出参：itemName、beginValue、endValue、changeValue、changeRate', 'fact_balance_sheet_snapshot'),
    ('获取资产负债表明细', 'GET', '/api/nongjing/finance/balance-sheet/detail', '返回资产负债表行次明细', '入参：regionCode、snapshotDate、statementType、pageNum、pageSize；出参：lineNo、itemName、beginValue、endValue', 'fact_balance_sheet_snapshot'),
    ('导出资产负债表明细', 'POST', '/api/nongjing/finance/balance-sheet/export', '导出资产负债表明细', '入参：snapshotDate、regionCode、statementType；出参：taskId、status', 'fact_balance_sheet_snapshot'),
    ('获取收益分配分析', 'GET', '/api/nongjing/finance/income-distribution/analysis', '返回收益分配结构及金额分析', '入参：regionCode、snapshotDate；出参：itemName、amount、rate', 'fact_income_distribution_snapshot'),
    ('获取资产一张图统计', 'GET', '/api/nongjing/assets/summary', '获取资产专题总额、经营性资产、非经营性资产及排行', '入参：regionCode、yearMonth；出参：assetTotal、businessAssetTotal、nonBusinessAssetTotal、rankList', 'fact_asset_summary_monthly、fact_asset_item'),
    ('获取资源一张图统计', 'GET', '/api/nongjing/resources/summary', '获取农用地、建设用地、未利用地等资源统计', '入参：regionCode、yearMonth；出参：farmLandArea、constructLandArea、unusedLandArea、rankList', 'fact_resource_summary_monthly、fact_resource_item'),
    ('获取合同一张图统计', 'GET', '/api/nongjing/contracts/summary', '获取合同数量、金额、应收应付统计', '入参：regionCode、yearMonth；出参：contractCount、contractAmount、receivableAmount、payableAmount', 'fact_contract_summary_monthly、fact_contract_main'),
    ('获取合同收付款明细', 'GET', '/api/nongjing/contracts/funds/list', '获取合同收款计划、付款计划和执行明细', '入参：regionCode、contractType、fundType、pageNum、pageSize；出参：contractNo、planDate、planAmount、execAmount、status', 'fact_contract_receipt_plan、fact_contract_payment_plan、fact_contract_fund_exec'),
    ('获取成员一张图统计', 'GET', '/api/nongjing/members/summary', '获取户数、人数、成员/非成员、男女分布等统计', '入参：regionCode、yearMonth；出参：householdCount、personCount、memberCount、nonMemberCount、maleCount、femaleCount', 'fact_member_summary_monthly、fact_member_household、fact_member_person'),
    ('获取股权一张图统计', 'GET', '/api/nongjing/equity/summary', '获取总股值、总分红、股权类型分布及排行', '入参：regionCode、yearMonth；出参：equityTotalValue、dividendTotalValue、typeDistribution、rankList', 'fact_equity_summary_monthly、fact_equity_holder'),
    ('获取股权分红明细', 'GET', '/api/nongjing/equity/dividends/list', '获取分红明细及分红对象信息', '入参：regionCode、yearMonth、pageNum、pageSize；出参：holderName、equityType、dividendAmount、dividendDate', 'fact_equity_dividend_detail'),
    ('获取预警规则列表', 'GET', '/api/nongjing/warnings/rules', '获取预警规则配置列表', '入参：moduleCode、status；出参：ruleId、ruleName、metricCode、compareOp、threshold、level', 'warning_rule'),
    ('保存预警规则', 'POST', '/api/nongjing/warnings/rules', '新增或修改预警规则', '入参：ruleName、metricCode、compareOp、threshold、level、scope；出参：ruleId、status', 'warning_rule'),
    ('获取预警事件列表', 'GET', '/api/nongjing/warnings/events', '分页查询预警事件', '入参：regionCode、yearMonth、level、handleStatus、pageNum、pageSize；出参：eventId、ruleName、triggerValue、level、handleStatus', 'fact_warning_event'),
    ('处置预警事件', 'POST', '/api/nongjing/warnings/events/{eventId}/handle', '登记预警处置结果', '入参：handleResult、handleUser、handleRemark；出参：eventId、handleStatus、handleTime', 'fact_warning_event、fact_warning_handle_log'),
    ('导出预警事件', 'POST', '/api/nongjing/warnings/events/export', '异步导出预警事件结果集', '入参：筛选快照、exportType；出参：taskId、status', 'fact_warning_event'),
]

village_interfaces = [
    ('获取账套列表', 'GET', '/api/village-finance/ledgers', '获取账套和启停状态列表', '入参：orgCode、status；出参：ledgerId、ledgerName、status', '账套表'),
    ('获取会计期间列表', 'GET', '/api/village-finance/periods', '获取会计期间及结账状态', '入参：ledgerId、year；出参：periodId、periodName、closeStatus', '会计期间表'),
    ('获取会计科目树', 'GET', '/api/village-finance/subjects/tree', '获取科目树和辅助核算配置', '入参：ledgerId；出参：subjectCode、subjectName、parentCode、assistFlag', '会计科目表、科目辅助核算配置表'),
    ('获取记账凭证列表', 'GET', '/api/village-finance/vouchers', '按编号、日期、状态查询凭证列表', '入参：ledgerId、periodId、voucherNo、auditStatus、pageNum、pageSize；出参：voucherId、voucherNo、voucherDate、totalAmount、auditStatus', '凭证主表'),
    ('获取记账凭证详情', 'GET', '/api/village-finance/vouchers/{voucherId}', '获取凭证头、分录和附件详情', '入参：voucherId；出参：凭证头、entries、attachments', '凭证主表、凭证明细表、凭证附件表'),
    ('新增记账凭证', 'POST', '/api/village-finance/vouchers', '新增手工凭证', '入参：ledgerId、periodId、summary、entries[]、attachmentIds[]；出参：voucherId、voucherNo、status', '凭证主表、凭证明细表、凭证附件表'),
    ('修改记账凭证', 'PUT', '/api/village-finance/vouchers/{voucherId}', '修改未审核凭证', '入参：voucherId、summary、entries[]；出参：voucherId、status', '凭证主表、凭证明细表'),
    ('审核记账凭证', 'POST', '/api/village-finance/vouchers/{voucherId}/audit', '审核通过凭证', '入参：voucherId、auditRemark；出参：voucherId、auditStatus、auditTime', '凭证主表、凭证审核日志表'),
    ('反审核记账凭证', 'POST', '/api/village-finance/vouchers/{voucherId}/unaudit', '反审核已审核凭证', '入参：voucherId、reason；出参：voucherId、auditStatus', '凭证主表、凭证审核日志表'),
    ('删除记账凭证', 'DELETE', '/api/village-finance/vouchers/{voucherId}', '删除或回收未生效凭证', '入参：voucherId；出参：voucherId、deletedFlag', '凭证回收站表'),
    ('上传凭证附件', 'POST', '/api/village-finance/vouchers/{voucherId}/attachments', '为凭证上传附件', '入参：file、fileName；出参：attachmentId、fileUrl', '凭证附件表'),
    ('导出记账凭证', 'POST', '/api/village-finance/vouchers/export', '导出凭证列表或凭证明细', '入参：筛选条件、exportType；出参：taskId、status', '凭证主表、凭证明细表'),
    ('获取银行日记账列表', 'GET', '/api/village-finance/journals/bank', '查询银行日记账', '入参：accountId、startDate、endDate、pageNum、pageSize；出参：txnDate、summary、incomeAmount、expenseAmount、balance', '银行日记账表'),
    ('新增银行日记账', 'POST', '/api/village-finance/journals/bank', '新增银行收支流水', '入参：accountId、txnDate、summary、amount、direction；出参：journalId、status', '银行日记账表'),
    ('获取现金日记账列表', 'GET', '/api/village-finance/journals/cash', '查询现金日记账', '入参：accountId、startDate、endDate、pageNum、pageSize；出参：txnDate、summary、incomeAmount、expenseAmount、balance', '现金日记账表'),
    ('新增现金日记账', 'POST', '/api/village-finance/journals/cash', '新增现金收支流水', '入参：accountId、txnDate、summary、amount、direction；出参：journalId、status', '现金日记账表'),
    ('获取账户汇总', 'GET', '/api/village-finance/funds/account-summary', '汇总银行账户和现金账户余额', '入参：ledgerId、periodId；出参：accountName、accountType、balance、lastTxnDate', '银行账户表、现金账户表、银行日记账表、现金日记账表'),
    ('获取收支类别汇总', 'GET', '/api/village-finance/funds/category-summary', '按收支类别统计金额和笔数', '入参：ledgerId、periodId、direction；出参：categoryCode、categoryName、amount、txnCount', '收入模板表、支出模板表、银行日记账表、现金日记账表'),
    ('获取现金账户列表', 'GET', '/api/village-finance/accounts/cash', '查询现金账户主档', '入参：ledgerId、status；出参：accountId、accountName、initBalance、currentBalance', '现金账户表'),
    ('获取银行账户列表', 'GET', '/api/village-finance/accounts/bank', '查询银行账户主档', '入参：ledgerId、status；出参：accountId、bankName、accountNo、currentBalance', '银行账户表'),
    ('获取收入模板列表', 'GET', '/api/village-finance/templates/income', '查询收入模板', '入参：ledgerId、keyword；出参：templateId、templateName、subjectCode、summaryRule', '收入模板表、收入模板明细表'),
    ('获取支出模板列表', 'GET', '/api/village-finance/templates/expense', '查询支出模板', '入参：ledgerId、keyword；出参：templateId、templateName、subjectCode、summaryRule', '支出模板表、支出模板明细表'),
    ('获取内部转账列表', 'GET', '/api/village-finance/internal-transfers', '查询内部转账记录', '入参：fromAccountId、toAccountId、startDate、endDate；出参：transferId、txnDate、amount、status', '内部转账表'),
    ('新增内部转账', 'POST', '/api/village-finance/internal-transfers', '发起内部转账并生成凭证', '入参：fromAccountId、toAccountId、amount、txnDate；出参：transferId、voucherId', '内部转账表、凭证主表'),
    ('获取对账结果', 'GET', '/api/village-finance/reconciliations', '查询出纳对账结果', '入参：ledgerId、periodId、accountType；出参：accountName、bookBalance、bankBalance、diffAmount', '出纳对账表'),
    ('获取未审核凭证列表', 'GET', '/api/village-finance/reconciliations/pending-vouchers', '查询未审核凭证用于对账', '入参：ledgerId、periodId；出参：voucherId、voucherNo、voucherDate、amount', '凭证主表'),
    ('获取序时账', 'GET', '/api/village-finance/books/journal', '查询序时账并支持穿透凭证', '入参：ledgerId、periodId、subjectCode、pageNum、pageSize；出参：txnDate、summary、subjectName、debitAmount、creditAmount、voucherId', '账簿日志表、凭证主表、凭证明细表'),
    ('获取科目余额表', 'GET', '/api/village-finance/books/subject-balance', '查询科目余额表', '入参：ledgerId、periodId、subjectLevel；出参：subjectCode、subjectName、beginBalance、debitAmount、creditAmount、endBalance', '科目余额表'),
    ('获取总账', 'GET', '/api/village-finance/books/general-ledger', '查询总账', '入参：ledgerId、periodId、subjectCode；出参：subjectCode、subjectName、beginBalance、periodAmount、endBalance', '总账快照/账簿类结果表'),
    ('获取明细账', 'GET', '/api/village-finance/books/detail-ledger', '查询明细账', '入参：ledgerId、periodId、subjectCode、pageNum、pageSize；出参：txnDate、summary、debitAmount、creditAmount、balance', '明细账结果集'),
    ('获取财务三表', 'GET', '/api/village-finance/reports/statements', '生成并查询资产负债表、利润表、现金流量表', '入参：ledgerId、periodId、statementType；出参：statementHeader、statementLines[]', '报表定义表、报表行次定义表、报表公式表、报表快照表'),
    ('导出报表', 'POST', '/api/village-finance/reports/export', '导出报表文件', '入参：statementType、periodId、exportType；出参：taskId、status', '报表快照表、报表导出记录表'),
    ('获取资产卡片列表', 'GET', '/api/village-finance/assets/cards', '查询资产卡片', '入参：assetType、status、pageNum、pageSize；出参：assetId、assetName、originValue、netValue、status', '资产卡片表'),
    ('新增资产卡片', 'POST', '/api/village-finance/assets/cards', '新增资产卡片', '入参：assetName、assetType、buyDate、originValue、deprYears；出参：assetId、status', '资产卡片表'),
    ('获取资产变更记录', 'GET', '/api/village-finance/assets/changes', '查询资产变更和折旧信息', '入参：assetId、changeType；出参：changeDate、changeType、beforeValue、afterValue', '资产变更表、资产折旧表'),
    ('获取资产盘点列表', 'GET', '/api/village-finance/assets/inventories', '查询资产盘点任务和结果', '入参：periodId、status；出参：inventoryId、assetCount、diffCount、status', '资产盘点表、资产盘点明细表'),
    ('获取合同列表', 'GET', '/api/village-finance/contracts', '查询合同信息', '入参：contractType、status、pageNum、pageSize；出参：contractId、contractNo、contractName、amount、status', '合同主表'),
    ('新增合同', 'POST', '/api/village-finance/contracts', '新增合同及附件', '入参：contractName、contractType、signDate、amount、partyA、partyB；出参：contractId、status', '合同主表、合同附件表'),
    ('获取合同收付款计划', 'GET', '/api/village-finance/contracts/{contractId}/payments', '查询合同收付款计划和执行', '入参：contractId；出参：planDate、planAmount、paidAmount、status', '合同收付款表'),
    ('获取审批待办列表', 'GET', '/api/village-finance/approvals/todos', '查询在线审批待办任务', '入参：userId、taskStatus、pageNum、pageSize；出参：taskId、bizType、bizNo、currentNode、taskStatus', '审批任务表、审批实例表'),
    ('审批支付单', 'POST', '/api/village-finance/payments/{paymentId}/approve', '审批支付单据', '入参：paymentId、approveResult、approveRemark；出参：paymentId、instanceStatus、nextNode', '支付单主表、支付单审批关联表、审批历史表'),
    ('获取银行流水', 'GET', '/api/village-finance/bank-link/transactions', '查询银农直联交易流水', '入参：accountId、startDate、endDate、pageNum、pageSize；出参：txnId、txnTime、amount、direction、status', '银行交易表、银行流水表'),
    ('获取银行余额', 'GET', '/api/village-finance/bank-link/balances', '查询银农直联账户余额', '入参：accountId；出参：accountId、availableBalance、bookBalance、syncTime', '银行余额表'),
    ('获取电子回单', 'GET', '/api/village-finance/bank-link/receipts', '查询银行电子回单', '入参：accountId、txnId、pageNum、pageSize；出参：receiptId、txnId、fileUrl、status', '银行回单表、电子回单文件表'),
    ('获取党组织列表', 'GET', '/api/village-finance/party/orgs', '查询党组织信息', '入参：orgCode、status；出参：partyOrgId、partyOrgName、secretaryName、status', '党组织表'),
    ('获取党员列表', 'GET', '/api/village-finance/party/members', '查询党员信息', '入参：partyOrgId、keyword、pageNum、pageSize；出参：memberId、memberName、joinDate、status', '党员表'),
    ('获取学习课程列表', 'GET', '/api/village-finance/party/courses', '查询在线学习课程', '入参：keyword、courseStatus；出参：courseId、courseName、teacherName、status', '在线学习课程表、课程章节表'),
    ('获取在线考试列表', 'GET', '/api/village-finance/party/exams', '查询考试和成绩', '入参：courseId、pageNum、pageSize；出参：examId、examName、scoreFlag、publishStatus', '在线考试表、考试题目表、考试成绩表'),
    ('生成结转试算', 'POST', '/api/village-finance/period-close/trial-balance', '生成结转前试算结果', '入参：ledgerId、periodId、templateId；出参：trialId、checkResult、errorCount', '期末结账校验表、损益结转主表'),
    ('执行月结', 'POST', '/api/village-finance/period-close/close', '执行月结并更新期间状态', '入参：ledgerId、periodId、closeType；出参：closeId、closeStatus、closeTime', '期末结账主表、会计期间表、期末处理日志表'),
    ('反结账', 'POST', '/api/village-finance/period-close/reopen', '对已结账期间执行反结账', '入参：ledgerId、periodId、reason；出参：reopenId、periodStatus、operateTime', '反结账记录表、会计期间表'),
]

common_query_example = '''示例接口：GET /api/nongjing/overview/cards?regionCode=460100&yearMonth=2026-06\n返回示例：{\n  "code": "0",\n  "message": "success",\n  "data": {\n    "countyCount": 18,\n    "townCount": 214,\n    "orgCount": 2689,\n    "unclosedCount": 126,\n    "noLedgerCount": 32\n  },\n  "requestId": "req-202606100001",\n  "timestamp": "2026-06-10 10:00:00"\n}'''

common_write_example = '''示例接口：POST /api/village-finance/vouchers\n请求示例：{\n  "ledgerId": "LD20260001",\n  "periodId": "2026-06",\n  "summary": "支付办公用品采购款",\n  "entries": [\n    {"subjectCode": "5601", "debitAmount": 5000.00, "creditAmount": 0},\n    {"subjectCode": "1002", "debitAmount": 0, "creditAmount": 5000.00}\n  ],\n  "attachmentIds": ["ATT0001", "ATT0002"]\n}\n响应示例：{\n  "code": "0",\n  "message": "success",\n  "data": {"voucherId": "V2026060001", "voucherNo": "记-2026-06-0001", "status": "DRAFT"},\n  "requestId": "req-202606100002",\n  "timestamp": "2026-06-10 10:05:00"\n}'''

error_codes = [
    ('IF0001', '参数校验失败', '必填字段缺失、格式错误或取值非法'),
    ('IF0002', '未登录或登录失效', 'Token 缺失、过期或校验失败'),
    ('IF0003', '无访问权限', '角色权限或数据权限不足'),
    ('IF0004', '对象不存在', '按主键或业务编码未查询到对象'),
    ('IF0005', '状态非法', '当前状态不允许执行提交、审核、反审核、删除等操作'),
    ('IF0006', '外部接口调用失败', '银农直联、短信、外部系统调用失败'),
    ('IF0007', '导出任务失败', '导出执行异常或文件生成失败'),
    ('IF0008', '附件处理失败', '上传、下载、预览或回单获取失败'),
]


def set_cn_font(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def add_run(para, text, size=12, bold=False, center=False):
    run = para.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run


def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_cn_font(run, 10)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    size_map = {1: 16, 2: 14, 3: 12, 4: 11}
    add_run(p, text, size=size_map.get(level, 12), bold=True)


def paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, text, size=12)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.line_spacing = 1.5
        add_run(p, item, size=12)


def interface_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['接口名称', '方法', 'URL', '用途', '关键参数与返回', '主要来源']
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, 9, True)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, 9, False)


def write_md(path, title, rows):
    lines = [f'# {title}', '', '| 接口名称 | 方法 | URL | 用途 | 关键参数与返回 | 主要来源 |', '|---|---|---|---|---|---|']
    for row in rows:
        safe = [str(x).replace('|', '\\|').replace('\n', '<br>') for x in row]
        lines.append('| ' + ' | '.join(safe) + ' |')
    path.write_text('\n'.join(lines), encoding='utf-8')

write_md(MD_NONGJING, '农经一张图接口明细（第一批）', nongjing_interfaces)
write_md(MD_VILLAGE, '村委财务事务管理系统接口明细（第一批）', village_interfaces)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
add_page_number(section)

add_run(doc.add_paragraph(), '海南数字“三农”金融服务平台', size=20, bold=True, center=True)
add_run(doc.add_paragraph(), '需求规格说明书（第三版）', size=18, bold=True, center=True)
add_run(doc.add_paragraph(), '适用范围：接口编写原则、公共接口规范、第一批详细接口正文', size=12, center=True)
add_run(doc.add_paragraph(), '版本：V3.0', size=12, center=True)
doc.add_paragraph('')
paragraph(doc, '说明：本版在第二版基础上，新增“农经一张图”和“村委财务事务管理系统”两大模块的第一批详细接口清单。其余 4 个模块保留目录骨架，下一轮继续补充。')

heading(doc, '6. 外部接口需求', 1)
paragraph(doc, '本章只保留接口编写依据、接口规范和接口文档内容，不重复写需求分析正文。')
heading(doc, '6.1 用户界面（UI）要求', 2)
numbered(doc, [
    '页面有展示必有查询接口，页面有操作必有写接口，页面有状态流转必有流程接口。',
    '图表、卡片、列表、树、地图和弹窗等界面元素均需在接口文档中明确数据来源。',
])
heading(doc, '6.2 硬件接口', 2)
paragraph(doc, '当前阶段未识别到专用硬件接口，若后续接入扫描、打印、读卡或物联网终端，另行补充。')
heading(doc, '6.3 软件接口', 2)
heading(doc, '6.3.1 接口编写原则说明', 3)
numbered(doc, [
    '接口编写以原型图设计要求为场景依据，以数据结构文档为字段依据，以映射文档为归属依据。',
    '接口章节只写接口定义，不回填需求分析正文。',
    '每个接口必须说明调用方、用途、方法、URL、关键入参、关键出参和主要来源。',
])
heading(doc, '6.3.2 公共接口规范', 3)
paragraph(doc, '统一返回结构：code、message、data、requestId、timestamp；分页统一采用 pageNum、pageSize、total、list；排序统一采用 orderBy、orderDirection。')
heading(doc, '6.3.2.1 标准查询接口示例', 4)
paragraph(doc, common_query_example)
heading(doc, '6.3.2.2 标准写入接口示例', 4)
paragraph(doc, common_write_example)
heading(doc, '6.3.2.3 通用错误码', 4)
err_table = doc.add_table(rows=1, cols=3)
err_table.alignment = WD_TABLE_ALIGNMENT.CENTER
err_table.style = 'Table Grid'
err_headers = ['错误码', '名称', '说明']
for i, h in enumerate(err_headers):
    err_table.rows[0].cells[i].text = h
for cell in err_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            set_cn_font(r, 10, True)
for code, name, desc in error_codes:
    row = err_table.add_row().cells
    row[0].text = code
    row[1].text = name
    row[2].text = desc
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, 10, False)

heading(doc, '6.3.3 农经一张图接口正文（第一批）', 3)
paragraph(doc, '接口来源：农经一张图原型需求梳理和数据结构文档。当前优先补齐首页总览、财务一张图、资产/资源/合同/成员/股权、预警监督等核心接口。')
interface_table(doc, nongjing_interfaces)

heading(doc, '6.3.4 村委财务事务管理系统接口正文（第一批）', 3)
paragraph(doc, '接口来源：村委财务事务管理系统原型需求梳理和数据结构文档。当前优先补齐财务中心、报表中心、资产管理、合同管理、在线审批、银农直联、基层党建、期末处理等核心接口。')
interface_table(doc, village_interfaces)

for idx, module in enumerate(modules[2:], start=5):
    heading(doc, f'6.3.{idx} {module["name"]}接口目录', 3)
    paragraph(doc, f'原型来源：{module["source_req"]}')
    paragraph(doc, f'数据来源：{module["source_data"]}')
    paragraph(doc, '当前保留接口目录骨架，后续按同一模板继续补充详细接口正文。')

heading(doc, '7. 系统部署与运行需求', 1)
paragraph(doc, PLACEHOLDER)

DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOC_OUT))
print(DOC_OUT)
print(MD_NONGJING)
print(MD_VILLAGE)
