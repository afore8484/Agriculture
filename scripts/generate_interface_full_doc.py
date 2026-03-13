from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r'E:/workspaces/Agriculture/prototype')
DOC_OUT = BASE / 'output' / 'doc' / 'spec_analysis_draft_v4_interfaces_full.docx'
MD_INCOME = BASE / 'docs' / 'interface_detail_income_ledger.md'
MD_PERF = BASE / 'docs' / 'interface_detail_performance.md'
MD_AUDIT = BASE / 'docs' / 'interface_detail_audit.md'
MD_ENTITY = BASE / 'docs' / 'interface_detail_entity_supervision.md'
MD_NONGJING = BASE / 'docs' / 'interface_detail_nongjing.md'
MD_VILLAGE = BASE / 'docs' / 'interface_detail_village_finance.md'
PLACEHOLDER = '【待补充】本章节暂按模板保留，后续继续补充。'


def parse_md_table(path: Path):
    rows = []
    lines = path.read_text(encoding='utf-8').splitlines()
    for line in lines:
        if not line.startswith('|'):
            continue
        if '---' in line or '接口名称' in line:
            continue
        parts = [p.strip() for p in line.strip().strip('|').split('|')]
        if len(parts) == 6:
            rows.append(tuple(parts))
    return rows

income_interfaces = [
    ('获取区划树', 'GET', '/api/income-ledger/common/regions', '获取五级区划树和层级信息', '入参：parentCode、level；出参：regionCode、regionName、regionLevel', 'adm_region'),
    ('获取收入记录列表', 'GET', '/api/income-ledger/records', '按区划、年度、季度查询台账记录', '入参：regionCode、year、quarter、status、pageNum、pageSize；出参：recordId、regionName、year、quarter、totalIncome、status', '经济收入台账主记录'),
    ('获取收入记录详情', 'GET', '/api/income-ledger/records/{recordId}', '查看台账主记录及收入明细', '入参：recordId；出参：主记录、明细项、附件列表', '经济收入台账主记录、经济收入台账明细、台账附件数据'),
    ('新增收入记录', 'POST', '/api/income-ledger/records', '新增台账记录', '入参：regionCode、year、quarter、incomeItems[]；出参：recordId、status', '经济收入台账主记录、经济收入台账明细'),
    ('修改收入记录', 'PUT', '/api/income-ledger/records/{recordId}', '修改草稿或退回记录', '入参：recordId、incomeItems[]、remark；出参：recordId、status', '经济收入台账主记录、经济收入台账明细'),
    ('删除收入记录', 'DELETE', '/api/income-ledger/records/{recordId}', '删除未锁定台账记录', '入参：recordId；出参：recordId、deletedFlag', '经济收入台账主记录'),
    ('导入收入记录', 'POST', '/api/income-ledger/records/import', '批量导入台账数据', '入参：file、templateVersion；出参：taskId、successCount、failCount', '导入导出任务数据、经济收入台账主记录'),
    ('导出收入记录', 'POST', '/api/income-ledger/records/export', '异步导出台账列表或明细', '入参：筛选条件、exportType；出参：taskId、status', '导入导出任务数据'),
    ('获取信息收集统计', 'GET', '/api/income-ledger/stats/collection', '查看填写率、上报率、进度统计', '入参：year、quarter、regionCode；出参：fillRate、submitRate、pendingCount、doneCount', '信息收集统计快照'),
    ('获取区划汇总表', 'GET', '/api/income-ledger/stats/region-summary', '查看按区划汇总的收入统计', '入参：year、quarter、regionLevel、regionCode；出参：regionName、totalIncome、totalExpense、netIncome', '区划汇总数据'),
    ('获取年度统计趋势', 'GET', '/api/income-ledger/stats/annual-trend', '查看年度收入趋势和对比', '入参：regionCode、startYear、endYear；出参：year、totalIncome、growthRate', '年度统计数据'),
    ('获取预警规则列表', 'GET', '/api/income-ledger/warnings/rules', '查看收入增长预警规则', '入参：regionCode、status；出参：ruleId、compareType、threshold、enabledFlag', '预警规则配置'),
    ('保存预警规则', 'POST', '/api/income-ledger/warnings/rules', '新增或修改预警规则', '入参：compareType、threshold、scope、needAttachment；出参：ruleId、status', '预警规则配置'),
    ('获取收入增长预警列表', 'GET', '/api/income-ledger/warnings/events', '分页查询预警事件', '入参：regionCode、year、quarter、handleStatus、pageNum、pageSize；出参：eventId、regionName、compareValue、threshold、handleStatus', '预警事件数据'),
    ('处理收入增长预警', 'POST', '/api/income-ledger/warnings/events/{eventId}/handle', '登记预警处置情况', '入参：handleResult、handleRemark、attachmentIds；出参：eventId、handleStatus', '预警事件数据、操作日志数据'),
    ('获取GDP目标列表', 'GET', '/api/income-ledger/gdp-targets', '查询区县年度GDP目标', '入参：regionCode、year、status；出参：targetId、targetValue、reportStatus', 'GDP增速目标记录'),
    ('保存GDP目标', 'POST', '/api/income-ledger/gdp-targets', '新增或修改GDP目标', '入参：regionCode、year、targetValue、remark；出参：targetId、status', 'GDP增速目标记录'),
    ('获取发展思路列表', 'GET', '/api/income-ledger/strategies', '查询发展思路和战略记录', '入参：regionCode、year；出参：strategyId、title、versionNo、status', '发展思路/战略记录'),
    ('保存发展思路', 'POST', '/api/income-ledger/strategies', '新增或修改发展思路记录', '入参：regionCode、year、title、content；出参：strategyId、versionNo', '发展思路/战略记录'),
    ('保存移动端草稿', 'POST', '/api/income-ledger/mobile/drafts', '保存手机端草稿', '入参：regionCode、year、quarter、draftJson；出参：draftId、status', '手机端提报数据'),
    ('提交移动端填报', 'POST', '/api/income-ledger/mobile/submissions', '提交手机端填报数据', '入参：draftId、submitRemark；出参：recordId、submitStatus', '手机端提报数据、上报流转记录'),
]

performance_interfaces = [
    ('获取评价事项列表', 'GET', '/api/performance/items', '查询评价事项主记录', '入参：year、status、regionCode、pageNum、pageSize；出参：itemId、itemName、periodName、status', '评价事项数据'),
    ('获取评价对象列表', 'GET', '/api/performance/objects', '查询参与评价的区划或组织对象', '入参：itemId、regionCode；出参：objectId、objectName、regionCode、ownerDept', '评价事项数据、组织机构主数据'),
    ('获取指标树', 'GET', '/api/performance/indicators/tree', '查询指标定义和层级结构', '入参：versionNo；出参：indicatorCode、indicatorName、unit、children', '指标定义主数据'),
    ('保存指标权重', 'POST', '/api/performance/indicators/weights', '配置指标权重与评分规则', '入参：versionNo、indicators[]；出参：versionNo、status', '评分规则与权重'),
    ('获取数据收集任务', 'GET', '/api/performance/collection/tasks', '查询数据收集任务和同步结果', '入参：itemId、taskStatus、pageNum、pageSize；出参：taskId、sourceSystem、executeStatus、executeTime', '数据收集任务数据'),
    ('新增采集数据', 'POST', '/api/performance/collection/records', '录入指标采集值', '入参：itemId、regionCode、indicatorCode、indicatorValue、sourceDesc；出参：recordId、status', '指标采集明细数据'),
    ('导入采集数据', 'POST', '/api/performance/collection/import', '批量导入采集明细', '入参：file、itemId；出参：taskId、successCount、failCount', '指标采集明细数据、数据收集任务数据'),
    ('获取待审核列表', 'GET', '/api/performance/reviews/todos', '查询待审核事项', '入参：reviewerId、status、pageNum、pageSize；出参：reviewId、itemName、regionName、currentStatus', '审核流转数据'),
    ('审核评价数据', 'POST', '/api/performance/reviews/{reviewId}/approve', '审核通过评价数据', '入参：reviewOpinion；出参：reviewId、status、nextNode', '审核流转数据'),
    ('退回评价数据', 'POST', '/api/performance/reviews/{reviewId}/reject', '退回评价数据', '入参：rejectReason；出参：reviewId、status', '审核流转数据'),
    ('触发评分计算', 'POST', '/api/performance/scores/calculate', '按版本规则计算得分', '入参：itemId、versionNo；出参：jobId、status', '评分规则与权重、评分结果数据'),
    ('获取评分结果列表', 'GET', '/api/performance/results', '查询总分、等级和排名', '入参：itemId、regionCode、pageNum、pageSize；出参：resultId、totalScore、rankNo、grade', '评分结果数据'),
    ('获取排名分析', 'GET', '/api/performance/results/rankings', '查看排名和分层分析', '入参：itemId、dimensionCode；出参：regionName、totalScore、rankNo、dimensionScore', '评分结果数据'),
    ('生成评价报告', 'POST', '/api/performance/reports/generate', '生成评价报告', '入参：itemId、templateId、regionCode；出参：reportId、status', '报告数据、报告模板主数据'),
    ('获取评价报告详情', 'GET', '/api/performance/reports/{reportId}', '查看报告正文与章节', '入参：reportId；出参：reportTitle、chapters、attachments', '报告数据'),
    ('导出评价报告', 'POST', '/api/performance/reports/export', '导出评价报告', '入参：reportId、exportType；出参：taskId、status', '报告数据'),
    ('获取统计分析', 'GET', '/api/performance/stats/summary', '查看区划汇总、趋势和对比', '入参：itemId、regionLevel、regionCode；出参：summaryCards、trendList、compareList', '评分结果数据、报告数据'),
]

audit_interfaces = [
    ('获取审计计划列表', 'GET', '/api/audit/plans', '查询审计计划', '入参：year、status、regionCode、pageNum、pageSize；出参：planId、planName、planYear、status', 'audit_plan'),
    ('新增审计计划', 'POST', '/api/audit/plans', '新增审计计划', '入参：planName、planYear、targetList[]；出参：planId、status', 'audit_plan、audit_plan_target'),
    ('修改审计计划', 'PUT', '/api/audit/plans/{planId}', '修改审计计划', '入参：planId、planName、targetList[]；出参：planId、status', 'audit_plan、audit_plan_change_log'),
    ('获取审计方案列表', 'GET', '/api/audit/schemes', '查询审计方案', '入参：planId、status、pageNum、pageSize；出参：schemeId、schemeName、status', 'audit_scheme'),
    ('新增审计方案', 'POST', '/api/audit/schemes', '新增审计方案和事项清单', '入参：planId、schemeName、items[]、members[]；出参：schemeId、status', 'audit_scheme、audit_scheme_item、audit_scheme_member'),
    ('审批审计方案', 'POST', '/api/audit/schemes/{schemeId}/approve', '审批审计方案', '入参：approveResult、approveRemark；出参：schemeId、status', 'audit_scheme_approval'),
    ('获取审计通知列表', 'GET', '/api/audit/notices', '查询审计通知', '入参：schemeId、status；出参：noticeId、noticeNo、status', 'audit_notice'),
    ('发送审计通知', 'POST', '/api/audit/notices', '发送审计通知并生成回执', '入参：schemeId、receiverList[]；出参：noticeId、status', 'audit_notice、audit_notice_receipt'),
    ('获取承诺书列表', 'GET', '/api/audit/commitments', '查询承诺书及流转记录', '入参：noticeId、status；出参：commitmentId、signStatus、signTime', 'audit_commitment、audit_commitment_log'),
    ('生成承诺书', 'POST', '/api/audit/commitments/generate', '生成审计承诺书', '入参：noticeId、templateId；出参：commitmentId、status', 'audit_commitment、audit_commitment_template'),
    ('获取资料调取单列表', 'GET', '/api/audit/material-requests', '查询资料调取单', '入参：schemeId、status、pageNum、pageSize；出参：requestId、requestNo、status', 'audit_material_request'),
    ('新增资料调取单', 'POST', '/api/audit/material-requests', '新增资料调取任务', '入参：schemeId、items[]、deadline；出参：requestId、status', 'audit_material_request、audit_material_request_item'),
    ('提交调取资料', 'POST', '/api/audit/material-submits', '提交调取资料和账套关联', '入参：requestId、files[]、accountBooks[]；出参：submitId、status', 'audit_material_submit、audit_material_accountbook、audit_material_data'),
    ('获取审计疑点列表', 'GET', '/api/audit/issues', '查询审计疑点', '入参：schemeId、riskLevel、status、pageNum、pageSize；出参：issueId、issueTitle、riskLevel、status', 'audit_issue'),
    ('查证审计疑点', 'POST', '/api/audit/issues/{issueId}/verify', '登记疑点查证结果', '入参：verifyResult、verifyRemark；出参：issueId、status', 'audit_issue_verify_log'),
    ('整改审计疑点', 'POST', '/api/audit/issues/{issueId}/rectify', '登记整改结果', '入参：rectifyResult、rectifyRemark；出参：issueId、status', 'audit_issue_rectify'),
    ('获取审计底稿列表', 'GET', '/api/audit/workpapers', '查询底稿列表', '入参：schemeId、status、pageNum、pageSize；出参：workpaperId、title、status', 'audit_workpaper'),
    ('提交底稿审核', 'POST', '/api/audit/workpapers/{workpaperId}/review', '提交底稿审核', '入参：reviewOpinion、reviewResult；出参：workpaperId、status', 'audit_workpaper_review'),
    ('生成审计报告草稿', 'POST', '/api/audit/reports/drafts/generate', '生成审计报告草稿', '入参：schemeId、issueIds[]；出参：draftId、status', 'audit_report_draft、audit_report_draft_item'),
    ('获取征求意见列表', 'GET', '/api/audit/opinions', '查询征求意见记录', '入参：draftId、status；出参：opinionId、receiverName、status', 'audit_opinion'),
    ('提交征求意见反馈', 'POST', '/api/audit/opinions/{opinionId}/feedback', '提交反馈意见', '入参：feedbackContent、attachments[]；出参：feedbackId、status', 'audit_opinion_feedback'),
    ('获取审计复核列表', 'GET', '/api/audit/reviews', '查询复核记录', '入参：draftId、status；出参：reviewId、reviewerName、status', 'audit_review、audit_review_item'),
    ('发布正式报告', 'POST', '/api/audit/reports/{reportId}/publish', '发布正式报告并归档', '入参：publishRemark；出参：reportId、status', 'audit_report、audit_report_archive'),
    ('生成审计决定书', 'POST', '/api/audit/decisions', '生成审计决定书', '入参：reportId、decisionContent；出参：decisionId、status', 'audit_decision'),
    ('获取审计统计报表', 'GET', '/api/audit/stats/summary', '查询统计快照和报表结果', '入参：year、regionCode；出参：cardStats、reportList', 'audit_stats_snapshot、rpt_report_instance'),
    ('导出审计报表', 'POST', '/api/audit/stats/export', '异步导出统计报表', '入参：reportCode、year、regionCode；出参：taskId、status', 'rpt_export_task'),
]

entity_interfaces = [
    ('获取主体端首页统计', 'GET', '/api/entity-platform/dashboard/subject', '查询主体端首页卡片和图表统计', '入参：entityId、yearMonth；出参：memberCount、goodsCount、capitalTotal、projectStats、chartData', 'rpt_entity_home_snapshot、entity_profile'),
    ('获取政府端大屏统计', 'GET', '/api/entity-platform/dashboard/gov', '查询政府端大屏统计', '入参：regionCode、yearMonth；出参：subjectCount、typeDistribution、fundDistribution、projectStats', 'rpt_gov_dashboard_stat、entity_profile'),
    ('获取金融端首页统计', 'GET', '/api/entity-platform/dashboard/bank', '查询金融端产品、申贷、放款、还款统计', '入参：orgCode、yearMonth；出参：productCount、loanApplyCount、loanAmount、repayCount', 'rpt_bank_home_snapshot、fin_loan_review_log'),
    ('获取经营主体列表', 'GET', '/api/entity-platform/entities', '查询经营主体档案', '入参：regionCode、subjectTypeCode、industryCode、pageNum、pageSize；出参：entityId、entityName、subjectTypeCode、status', 'entity_profile'),
    ('获取经营主体详情', 'GET', '/api/entity-platform/entities/{entityId}', '查询主体详情', '入参：entityId；出参：基础信息、经营信息、联系人、财务摘要', 'entity_profile'),
    ('获取报表报送记录', 'GET', '/api/entity-platform/entities/{entityId}/reports', '查询主体报表报送记录', '入参：entityId、reportMonth、reportStatus；出参：reportId、reportMonth、totalIncome、netProfit、reportStatus', 'entity_report_monthly、entity_report_record'),
    ('提交主体月报', 'POST', '/api/entity-platform/entities/{entityId}/reports', '提交主体月报及附件', '入参：reportMonth、totalIncome、totalCost、netProfit、attachmentIds[]；出参：reportId、reportStatus', 'entity_report_monthly、entity_report_file'),
    ('审核主体月报', 'POST', '/api/entity-platform/entity-reports/{reportId}/review', '审核主体月报', '入参：reviewResult、reviewRemark；出参：reportId、reportStatus', 'entity_report_review_log'),
    ('获取成员账户列表', 'GET', '/api/entity-platform/member-capitals', '查询成员出资和权益记录', '入参：entityId、memberName、pageNum、pageSize；出参：memberId、memberName、capitalAmount、capitalType', 'member_capital_txn'),
    ('新增成员出资', 'POST', '/api/entity-platform/member-capitals', '新增成员出资记录', '入参：entityId、memberName、capitalType、capitalSource、amount、capitalDate；出参：capitalId、status', 'member_capital_txn'),
    ('转增公积金', 'POST', '/api/entity-platform/member-capitals/bonus', '登记公积金转增', '入参：entityId、memberId、bonusAmount、bonusDate；出参：recordId、status', 'member_capital_bonus'),
    ('成员出资转让', 'POST', '/api/entity-platform/member-capitals/transfer', '登记出资转让', '入参：fromMemberId、toMemberId、transferAmount、transferDate；出参：transferId、status', 'member_capital_transfer'),
    ('成员出资继承', 'POST', '/api/entity-platform/member-capitals/inherit', '登记出资继承', '入参：originMemberId、targetMemberId、inheritAmount、inheritDate；出参：inheritId、status', 'member_capital_inherit'),
    ('获取成员交易记录', 'GET', '/api/entity-platform/member-trades', '查询成员交易记录', '入参：entityId、memberId、pageNum、pageSize；出参：tradeId、tradeDate、tradeType、amount', '成员交易记录表'),
    ('新增成员交易记录', 'POST', '/api/entity-platform/member-trades', '新增成员交易记录', '入参：entityId、memberId、tradeType、amount、tradeDate；出参：tradeId、status', '成员交易记录表'),
    ('获取盈余量化记录', 'GET', '/api/entity-platform/surplus/quantizations', '查询财政补助、捐赠、公积金量化记录', '入参：entityId、year、type；出参：recordId、type、amount、bizDate', 'surplus_subsidy_detail、surplus_donation_detail、surplus_reserve_fund_detail'),
    ('保存盈余分配', 'POST', '/api/entity-platform/surplus/distributions', '保存盈余分配方案', '入参：entityId、year、distributableAmount、returnRate、memberItems[]；出参：distributionId、status', '盈余分配主表/明细表'),
    ('获取项目列表', 'GET', '/api/entity-platform/projects', '查询项目发布和申报状态', '入参：regionCode、projectStatus、pageNum、pageSize；出参：projectId、projectName、supportAmount、applyStatus', '项目发布主表、项目申报主表'),
    ('发起项目申报', 'POST', '/api/entity-platform/projects/{projectId}/apply', '对项目发起申报', '入参：entityId、applyRemark、attachmentIds[]；出参：applyId、status', '项目申报主表、项目申报附件表'),
    ('获取示范申报列表', 'GET', '/api/entity-platform/demonstrations', '查询示范申报和监测记录', '入参：entityId、status、pageNum、pageSize；出参：demoId、demoLevel、applyStatus、monitorStatus', '示范申报表、示范监测表'),
    ('提交示范申报', 'POST', '/api/entity-platform/demonstrations', '提交示范申报', '入参：entityId、demoLevel、attachmentIds[]；出参：demoId、status', '示范申报表'),
    ('审核示范申报', 'POST', '/api/entity-platform/demonstrations/{demoId}/review', '审核示范申报', '入参：reviewResult、reviewRemark；出参：demoId、status', '示范申报审核表'),
    ('获取商品列表', 'GET', '/api/entity-platform/products', '查询商品主数据', '入参：entityId、keyword、pageNum、pageSize；出参：productId、productName、specification、unitCode', '商品主表'),
    ('获取库存汇总', 'GET', '/api/entity-platform/inventory/stocks', '查询库存汇总和期初库存', '入参：entityId、warehouseId、keyword；出参：productId、productName、stockQty、warehouseName', 'inv_stock_init、库存汇总表'),
    ('新增采购入库单', 'POST', '/api/entity-platform/inventory/purchase-in', '新增采购入库单', '入参：entityId、warehouseId、supplierId、items[]；出参：orderId、status', 'inv_purchase_in'),
    ('新增销售出库单', 'POST', '/api/entity-platform/inventory/sale-out', '新增销售出库单', '入参：entityId、warehouseId、customerId、items[]；出参：orderId、status', 'inv_sale_out'),
    ('新增库存调拨单', 'POST', '/api/entity-platform/inventory/transfers', '新增库存调拨单', '入参：entityId、fromWarehouseId、toWarehouseId、items[]；出参：transferId、status', 'inv_transfer_order'),
    ('获取金融产品列表', 'GET', '/api/entity-platform/finance/products', '查询金融产品', '入参：orgCode、productType、pageNum、pageSize；出参：productId、productName、rateType、loanTerm', '金融产品主表、fin_product_publish_log'),
    ('发起贷款申请', 'POST', '/api/entity-platform/loans/applications', '主体发起贷款申请', '入参：entityId、productId、applyAmount、applyTerm、attachmentIds[]；出参：applicationId、status', '贷款申请主表'),
    ('审核贷款申请', 'POST', '/api/entity-platform/loans/applications/{applicationId}/review', '银行审核贷款申请', '入参：reviewResult、reviewRemark、approvedAmount；出参：applicationId、status', 'fin_loan_review_log'),
    ('获取农产品行情', 'GET', '/api/entity-platform/market-prices', '查询农产品行情', '入参：productName、marketDate、pageNum、pageSize；出参：productName、highestPrice、lowestPrice、referencePrice', 'market_price_main、market_price_source'),
    ('同步农产品行情', 'POST', '/api/entity-platform/market-prices/sync', '手工触发行情同步', '入参：sourceCode、marketDate；出参：syncBatchId、status', 'market_price_sync_log'),
    ('获取通知公告列表', 'GET', '/api/entity-platform/notices', '查询通知公告', '入参：noticeType、keyword、pageNum、pageSize；出参：noticeId、title、publishTime、status', '通知公告主表'),
    ('获取政策法规列表', 'GET', '/api/entity-platform/policies', '查询政策法规', '入参：policyType、keyword、pageNum、pageSize；出参：policyId、title、publishDept、publishDate', '政策法规主表'),
]


def set_cn_font(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def add_run(para, text, size=12, bold=False, center=False):
    run = para.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run


def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_cn_font(run, 10)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    size_map = {1: 16, 2: 14, 3: 12}
    add_run(p, text, size=size_map.get(level, 12), bold=True)


def paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, text, size=12)


def interface_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['接口名称', '方法', 'URL', '用途', '关键参数与返回', '主要来源']
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, 9, True)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, 9, False)


def write_md(path: Path, title: str, rows):
    lines = [f'# {title}', '', '| 接口名称 | 方法 | URL | 用途 | 关键参数与返回 | 主要来源 |', '|---|---|---|---|---|---|']
    for row in rows:
        safe = [str(x).replace('|', '\\|').replace('\n', '<br>') for x in row]
        lines.append('| ' + ' | '.join(safe) + ' |')
    path.write_text('\n'.join(lines), encoding='utf-8')

# read previous module details
nongjing_rows = parse_md_table(MD_NONGJING)
village_rows = parse_md_table(MD_VILLAGE)

write_md(MD_INCOME, '村集体经济收入台账接口明细（第一版）', income_interfaces)
write_md(MD_PERF, '村集体绩效评价管理系统接口明细（第一版）', performance_interfaces)
write_md(MD_AUDIT, '村集体智能审计系统接口明细（第一版）', audit_interfaces)
write_md(MD_ENTITY, '新型农业经营主体财务经营监管系统接口明细（第一版）', entity_interfaces)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
add_page_number(section)

add_run(doc.add_paragraph(), '海南数字“三农”金融服务平台', size=20, bold=True, center=True)
add_run(doc.add_paragraph(), '需求规格说明书（第四版）', size=18, bold=True, center=True)
add_run(doc.add_paragraph(), '适用范围：第6章接口章节完整版', size=12, center=True)
add_run(doc.add_paragraph(), '版本：V4.0', size=12, center=True)
doc.add_paragraph('')
paragraph(doc, '说明：本版在第三版基础上补齐剩余 4 个模块的详细接口清单，形成第6章接口章节完整版。文档仍按你的要求只写接口，不回填需求分析正文。')

heading(doc, '6. 外部接口需求', 1)
paragraph(doc, '本章为接口章节完整版，覆盖公共接口规范以及 6 个业务模块的第一版详细接口清单。')
heading(doc, '6.1 用户界面（UI）要求', 2)
paragraph(doc, 'UI章节仅保留页面与接口的映射要求，不展开需求描述。所有页面展示、操作、状态流转、导入导出、附件上传都必须映射到接口。')
heading(doc, '6.2 硬件接口', 2)
paragraph(doc, '当前阶段未识别到专用硬件控制接口，后续如接入扫描、读卡、打印或物联网终端，再单独补充。')
heading(doc, '6.3 软件接口', 2)
paragraph(doc, '所有接口统一遵循第二版和第三版中已定义的接口编写原则、统一返回结构、统一错误码和统一分页规则。以下按模块输出详细接口清单。')

full_modules = [
    ('6.3.1 农经一张图接口正文', '接口来源：农经一张图原型需求梳理和数据结构文档。', nongjing_rows),
    ('6.3.2 村委财务事务管理系统接口正文', '接口来源：村委财务事务管理系统原型需求梳理和数据结构文档。', village_rows),
    ('6.3.3 村集体经济收入台账接口正文', '接口来源：村集体经济收入台账原型需求梳理和数据结构文档。', income_interfaces),
    ('6.3.4 村集体绩效评价管理系统接口正文', '接口来源：村集体绩效评价管理系统原型需求梳理和数据结构文档。', performance_interfaces),
    ('6.3.5 村集体智能审计系统接口正文', '接口来源：村集体智能审计系统原型需求梳理和数据结构文档。', audit_interfaces),
    ('6.3.6 新型农业经营主体财务经营监管系统接口正文', '接口来源：新型农业经营主体财务经营监管系统原型需求梳理和数据结构文档。', entity_interfaces),
]

for title, desc, rows in full_modules:
    heading(doc, title, 3)
    paragraph(doc, desc)
    interface_table(doc, rows)

heading(doc, '7. 系统部署与运行需求', 1)
paragraph(doc, PLACEHOLDER)

DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOC_OUT))
print(DOC_OUT)
print(MD_INCOME)
print(MD_PERF)
print(MD_AUDIT)
print(MD_ENTITY)
