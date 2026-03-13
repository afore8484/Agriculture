from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r'E:/workspaces/Agriculture/prototype')
DOC_OUT = BASE / 'output' / 'doc' / 'spec_analysis_draft_v5_interfaces_rechecked.docx'
REVIEW_MD = BASE / 'docs' / 'interface_review_notes.md'
COMMON_MD = BASE / 'docs' / 'interface_detail_common.md'
FILES = {
    'nongjing': BASE / 'docs' / 'interface_detail_nongjing.md',
    'village': BASE / 'docs' / 'interface_detail_village_finance.md',
    'income': BASE / 'docs' / 'interface_detail_income_ledger.md',
    'performance': BASE / 'docs' / 'interface_detail_performance.md',
    'audit': BASE / 'docs' / 'interface_detail_audit.md',
    'entity': BASE / 'docs' / 'interface_detail_entity_supervision.md',
}


def parse_md_table(path: Path):
    rows = []
    for line in path.read_text(encoding='utf-8').splitlines():
        if not line.startswith('|') or '---' in line or '接口名称' in line:
            continue
        parts = [p.strip() for p in line.strip().strip('|').split('|')]
        if len(parts) == 6:
            rows.append(tuple(parts))
    return rows


def write_md(path: Path, title: str, rows):
    lines = [f'# {title}', '', '| 接口名称 | 方法 | URL | 用途 | 关键参数与返回 | 主要来源 |', '|---|---|---|---|---|---|']
    for row in rows:
        safe = [str(x).replace('|', '\\|').replace('\n', '<br>') for x in row]
        lines.append('| ' + ' | '.join(safe) + ' |')
    path.write_text('\n'.join(lines), encoding='utf-8')

existing = {k: parse_md_table(v) for k, v in FILES.items()}

common_rows = [
    ('用户登录', 'POST', '/api/common/auth/login', '账号登录并返回访问令牌', '入参：username、password、captcha；出参：accessToken、refreshToken、userInfo', 'sys_user、sys_role、sys_user_role'),
    ('用户退出', 'POST', '/api/common/auth/logout', '注销当前登录会话', '入参：accessToken；出参：successFlag', '登录会话表/缓存'),
    ('获取当前用户信息', 'GET', '/api/common/auth/current-user', '获取当前用户和权限摘要', '入参：无；出参：userId、realName、roles、dataScope', 'sys_user、sys_role、sys_user_role'),
    ('获取角色权限树', 'GET', '/api/common/auth/permissions', '获取菜单和按钮权限', '入参：roleId、terminalType；出参：menuTree、permissionCodes', 'sys_permission、sys_role_permission/sys_role_menu'),
    ('获取组织树', 'GET', '/api/common/orgs/tree', '获取组织机构树', '入参：orgType、parentOrgCode；出参：orgCode、orgName、children', 'sys_org'),
    ('获取区划树', 'GET', '/api/common/regions/tree', '获取行政区划树', '入参：level、parentCode；出参：regionCode、regionName、children', 'adm_region/sys_region/dim_region'),
    ('获取字典项', 'GET', '/api/common/dicts/items', '获取通用字典项', '入参：dictType；出参：dictCode、dictName、dictValue', 'sys_dict_item/dim_dict_item'),
    ('上传附件', 'POST', '/api/common/files/upload', '上传附件并返回文件元数据', '入参：file、bizType；出参：fileId、fileName、fileUrl', 'sys_file_object、sys_attachment、对象存储'),
    ('下载附件', 'GET', '/api/common/files/{fileId}', '下载或预览附件', '入参：fileId；出参：fileUrl、contentType、fileName', 'sys_file_object、sys_attachment、对象存储'),
    ('查询导出任务', 'GET', '/api/common/export/tasks', '查询导出任务状态和结果文件', '入参：taskId、bizType；出参：taskId、status、fileUrl、failReason', 'rpt_export_task/导入导出任务数据'),
    ('查询操作日志', 'GET', '/api/common/logs/operations', '按模块查询操作日志', '入参：moduleCode、bizId、startTime、endTime；出参：logId、operatorName、actionType、operateTime', 'sys_log_operation/sys_log_export_task'),
    ('查询同步日志', 'GET', '/api/common/logs/sync', '按对接类型查询同步日志', '入参：syncType、bizDate、status；出参：syncBatchId、syncType、status、message', '接口同步日志/market_price_sync_log'),
]

village_add = [
    ('新增现金账户', 'POST', '/api/village-finance/accounts/cash', '新增现金账户主档', '入参：ledgerId、accountName、initBalance；出参：accountId、status', '现金账户表'),
    ('新增银行账户', 'POST', '/api/village-finance/accounts/bank', '新增银行账户主档', '入参：ledgerId、bankName、accountNo、initBalance；出参：accountId、status', '银行账户表'),
    ('保存收入模板', 'POST', '/api/village-finance/templates/income', '新增或修改收入模板', '入参：templateName、subjectCode、summaryRule；出参：templateId、status', '收入模板表、收入模板明细表'),
    ('保存支出模板', 'POST', '/api/village-finance/templates/expense', '新增或修改支出模板', '入参：templateName、subjectCode、summaryRule；出参：templateId、status', '支出模板表、支出模板明细表'),
    ('删除内部转账', 'DELETE', '/api/village-finance/internal-transfers/{transferId}', '删除错误内部转账记录', '入参：transferId；出参：transferId、deletedFlag', '内部转账表'),
    ('导出内部转账', 'POST', '/api/village-finance/internal-transfers/export', '导出内部转账记录', '入参：startDate、endDate、accountId；出参：taskId、status', '内部转账表'),
    ('获取初始化模板列表', 'GET', '/api/village-finance/init-templates', '查询出纳初始化模板', '入参：ledgerId、keyword；出参：templateId、templateName、versionNo', '出纳初始化模板表'),
    ('保存初始化模板', 'POST', '/api/village-finance/init-templates', '新增或修改初始化模板', '入参：templateName、versionNo、templateJson；出参：templateId、status', '出纳初始化模板表、出纳初始化模板明细表'),
    ('获取回收站凭证列表', 'GET', '/api/village-finance/vouchers/recycle-bin', '查询已删除凭证', '入参：ledgerId、pageNum、pageSize；出参：voucherId、voucherNo、deletedTime', '凭证回收站表'),
    ('还原回收站凭证', 'POST', '/api/village-finance/vouchers/{voucherId}/restore', '还原回收站凭证', '入参：voucherId；出参：voucherId、status', '凭证回收站表'),
    ('彻底删除回收站凭证', 'DELETE', '/api/village-finance/vouchers/recycle-bin/{voucherId}', '彻底删除回收站凭证', '入参：voucherId；出参：voucherId、deletedFlag', '凭证回收站表'),
    ('获取打印设置', 'GET', '/api/village-finance/print-settings', '查询凭证或报表打印设置', '入参：bizType；出参：paperSize、direction、marginConfig', '凭证打印配置表、报表打印配置表'),
    ('保存打印设置', 'POST', '/api/village-finance/print-settings', '保存打印设置', '入参：bizType、paperSize、direction、marginConfig；出参：configId、status', '凭证打印配置表、报表打印配置表'),
]

income_add = [
    ('获取特殊区域标签', 'GET', '/api/income-ledger/regions/tags', '查询特殊区域标签', '入参：regionCode、tagType；出参：tagId、regionCode、tagType、tagName', 'region_special_tag'),
    ('保存特殊区域标签', 'POST', '/api/income-ledger/regions/tags', '新增或修改特殊区域标签', '入参：regionCode、tagType、tagRemark；出参：tagId、status', 'region_special_tag'),
    ('获取5万/10万清单', 'GET', '/api/income-ledger/key-lists', '查询重点清单', '入参：listType、year、pageNum、pageSize；出参：listId、regionName、listType、status', '5万/10万清单配置'),
    ('导入5万/10万清单', 'POST', '/api/income-ledger/key-lists/import', '导入重点清单', '入参：listType、file；出参：taskId、status', '5万/10万清单配置、导入导出任务数据'),
    ('获取填报窗口', 'GET', '/api/income-ledger/report-windows', '查询填报时间窗口', '入参：bizType、year；出参：windowId、startTime、endTime、enabledFlag', '填报时间配置'),
    ('保存填报窗口', 'POST', '/api/income-ledger/report-windows', '新增或修改填报窗口', '入参：bizType、year、startTime、endTime、enabledFlag；出参：windowId、status', '填报时间配置'),
    ('获取上报流转记录', 'GET', '/api/income-ledger/submission-logs', '查询提交、退回、审核流转记录', '入参：recordId；出参：actionType、actionUser、actionTime、actionRemark', '上报流转记录'),
    ('上传台账附件', 'POST', '/api/income-ledger/records/{recordId}/attachments', '上传台账附件', '入参：recordId、file；出参：attachmentId、fileUrl', '台账附件数据'),
    ('查询同步日志', 'GET', '/api/income-ledger/sync-logs', '查询与三资系统的同步日志', '入参：syncType、bizDate、status；出参：syncBatchId、status、message', '接口同步日志'),
]

performance_add = [
    ('获取用户列表', 'GET', '/api/performance/system/users', '查询政府端用户', '入参：regionCode、accountLevel、pageNum、pageSize；出参：userId、username、realName、status', 'sys_user'),
    ('获取角色列表', 'GET', '/api/performance/system/roles', '查询角色和权限摘要', '入参：status；出参：roleId、roleCode、roleName、status', 'sys_role'),
    ('保存角色权限', 'POST', '/api/performance/system/roles/permissions', '配置角色权限', '入参：roleId、permissionIds[]；出参：roleId、status', 'sys_role_permission、sys_permission'),
    ('获取密码策略', 'GET', '/api/performance/system/password-policy', '查询密码策略', '入参：policyCode；出参：minLength、requireUppercase、requireDigit、passwordValidDays', 'sys_password_policy'),
    ('保存密码策略', 'POST', '/api/performance/system/password-policy', '保存密码策略', '入参：policyCode、minLength、requireUppercase、requireDigit；出参：policyId、status', 'sys_password_policy'),
    ('获取同步日志', 'GET', '/api/performance/system/sync-logs', '查询数据收集同步日志', '入参：taskId、status、pageNum、pageSize；出参：logId、sourceSystem、executeStatus、executeTime', '数据收集任务数据、日志与审计数据'),
]

audit_add = [
    ('获取审计事项库', 'GET', '/api/audit/base/items', '查询审计事项库', '入参：itemType、riskLevel、pageNum、pageSize；出参：itemId、itemCode、itemName、riskLevel', 'base_audit_item_lib'),
    ('保存审计事项', 'POST', '/api/audit/base/items', '新增或修改审计事项', '入参：itemCode、itemName、riskLevel、checkMethod；出参：itemId、status', 'base_audit_item_lib'),
    ('获取流程模板列表', 'GET', '/api/audit/base/process-templates', '查询流程模板', '入参：flowCode、status；出参：templateId、flowCode、flowName、status', 'base_process_template'),
    ('保存流程模板', 'POST', '/api/audit/base/process-templates', '新增或修改流程模板', '入参：flowCode、flowName、nodesJson；出参：templateId、status', 'base_process_template'),
    ('上传审计附件', 'POST', '/api/audit/common/attachments', '上传审计通用附件', '入参：bizType、bizId、file；出参：attachmentId、fileUrl', 'sys_attachment、sys_file_object'),
    ('获取流程实例详情', 'GET', '/api/audit/workflows/{instanceId}', '查询流程实例和任务节点', '入参：instanceId；出参：instanceId、currentNode、taskList[]', 'wf_instance、wf_task'),
]

entity_add = [
    ('获取店铺列表', 'GET', '/api/entity-platform/stores', '查询店铺信息', '入参：entityId、keyword、pageNum、pageSize；出参：storeId、storeName、status', '店铺主表'),
    ('获取客户列表', 'GET', '/api/entity-platform/customers', '查询客户信息', '入参：entityId、keyword、pageNum、pageSize；出参：customerId、customerName、customerType', '客户主表'),
    ('获取仓库列表', 'GET', '/api/entity-platform/warehouses', '查询仓库信息', '入参：entityId、warehouseType；出参：warehouseId、warehouseName、status', '仓库主表'),
    ('获取计量单位列表', 'GET', '/api/entity-platform/units', '查询计量单位配置', '入参：entityId、keyword；出参：unitId、unitName、convertRate', 'inv_unit_config'),
    ('新增采购退货单', 'POST', '/api/entity-platform/inventory/purchase-return', '新增采购退货单', '入参：entityId、supplierId、items[]；出参：orderId、status', 'inv_purchase_return'),
    ('新增销售退货单', 'POST', '/api/entity-platform/inventory/sale-return', '新增销售退货单', '入参：entityId、customerId、items[]；出参：orderId、status', 'inv_sale_return'),
    ('获取库存调拨单列表', 'GET', '/api/entity-platform/inventory/transfers', '查询库存调拨单', '入参：entityId、status、pageNum、pageSize；出参：transferId、fromWarehouseId、toWarehouseId、status', 'inv_transfer_order'),
    ('获取助农服务目录', 'GET', '/api/entity-platform/agri-services', '查询助农服务目录', '入参：serviceType、keyword、pageNum、pageSize；出参：serviceId、serviceName、serviceType', 'agri_service_item'),
    ('发布金融产品', 'POST', '/api/entity-platform/finance/products', '银行端发布金融产品', '入参：productName、productType、rateType、loanTerm；出参：productId、status', '金融产品主表、fin_product_publish_log'),
    ('记录放款结果', 'POST', '/api/entity-platform/loans/applications/{applicationId}/disburse', '登记放款结果', '入参：approvedAmount、disburseDate；出参：applicationId、status', '贷款放款记录表'),
    ('记录还款结果', 'POST', '/api/entity-platform/loans/applications/{applicationId}/repay', '登记还款结果', '入参：repayAmount、repayDate；出参：applicationId、status', '贷款还款记录表'),
    ('记录催收结果', 'POST', '/api/entity-platform/loans/applications/{applicationId}/collections', '登记催收结果', '入参：collectionType、collectionRemark、nextActionDate；出参：collectionId、status', 'fin_collection_record'),
    ('发布通知公告', 'POST', '/api/entity-platform/notices', '新增或发布通知公告', '入参：noticeType、title、content、publishTime；出参：noticeId、status', '通知公告主表'),
]

review_notes = '''# 接口复查结果与修正说明

## 结论
原 `173` 个接口可以覆盖“核心流程第一版”，但不能直接判定为已经完全满足需求文档和数据结构。复查后识别出 6 类高风险缺口，并已在 `v5` 中补齐。

## 已识别并修正的主要问题
1. 缺少公共接口定义：登录、权限、区划树、字典、附件、导出任务、同步日志未作为正式接口列出。
2. 村委财务事务管理系统缺少模板管理、初始化模板、回收站、打印设置等接口。
3. 村集体经济收入台账缺少特殊区域标签、5万/10万清单、填报窗口、附件与同步日志接口。
4. 村集体绩效评价管理系统缺少用户、角色、权限、密码策略等系统管理接口。
5. 村集体智能审计系统缺少审计事项库、流程模板、通用附件、流程实例接口。
6. 新型农业经营主体财务经营监管系统缺少店铺、客户、仓库、计量单位、退货、助农服务、放款/还款/催收、公告发布等接口。
'''

merged = {
    'common': common_rows,
    'nongjing': existing['nongjing'],
    'village': existing['village'] + village_add,
    'income': existing['income'] + income_add,
    'performance': existing['performance'] + performance_add,
    'audit': existing['audit'] + audit_add,
    'entity': existing['entity'] + entity_add,
}

write_md(COMMON_MD, '公共接口明细（第二版）', merged['common'])
write_md(FILES['village'], '村委财务事务管理系统接口明细（第二版）', merged['village'])
write_md(FILES['income'], '村集体经济收入台账接口明细（第二版）', merged['income'])
write_md(FILES['performance'], '村集体绩效评价管理系统接口明细（第二版）', merged['performance'])
write_md(FILES['audit'], '村集体智能审计系统接口明细（第二版）', merged['audit'])
write_md(FILES['entity'], '新型农业经营主体财务经营监管系统接口明细（第二版）', merged['entity'])
REVIEW_MD.write_text(review_notes, encoding='utf-8')


def set_cn_font(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def add_run(para, text, size=12, bold=False, center=False):
    run = para.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run


def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_cn_font(run, 10)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    size_map = {1: 16, 2: 14, 3: 12}
    add_run(p, text, size=size_map.get(level, 12), bold=True)


def paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, text, size=12)


def interface_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['接口名称', '方法', 'URL', '用途', '关键参数与返回', '主要来源']
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, 9, True)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, 9, False)

count_total = sum(len(v) for v in merged.values())

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
add_page_number(section)

add_run(doc.add_paragraph(), '海南数字“三农”金融服务平台', size=20, bold=True, center=True)
add_run(doc.add_paragraph(), '需求规格说明书（第五版）', size=18, bold=True, center=True)
add_run(doc.add_paragraph(), '适用范围：第6章接口章节复查修正版', size=12, center=True)
add_run(doc.add_paragraph(), '版本：V5.0', size=12, center=True)
doc.add_paragraph('')
paragraph(doc, f'说明：本版基于对原 173 个接口的复查结果进行了修正和补充，新增公共接口与模块缺失接口，当前接口总数为 {count_total}。')

heading(doc, '6. 外部接口需求', 1)
paragraph(doc, '本版为接口章节复查修正版。已按需求文档动作与数据结构表覆盖关系，对第一版接口进行补漏和修正。')
heading(doc, '6.1 复查结论', 2)
paragraph(doc, '复查结果表明，原 173 个接口已覆盖核心流程，但尚未完全满足需求文档和数据结构的全部接口场景。已在本版补齐公共接口、系统管理接口、模板与回收站接口、窗口与标签接口、事项库与流程模板接口，以及主体监管的退货/仓储/放款还款等接口。')
heading(doc, '6.2 公共接口', 2)
interface_table(doc, merged['common'])
heading(doc, '6.3 农经一张图接口正文', 2)
interface_table(doc, merged['nongjing'])
heading(doc, '6.4 村委财务事务管理系统接口正文', 2)
interface_table(doc, merged['village'])
heading(doc, '6.5 村集体经济收入台账接口正文', 2)
interface_table(doc, merged['income'])
heading(doc, '6.6 村集体绩效评价管理系统接口正文', 2)
interface_table(doc, merged['performance'])
heading(doc, '6.7 村集体智能审计系统接口正文', 2)
interface_table(doc, merged['audit'])
heading(doc, '6.8 新型农业经营主体财务经营监管系统接口正文', 2)
interface_table(doc, merged['entity'])
heading(doc, '7. 系统部署与运行需求', 1)
paragraph(doc, '【待补充】本章节暂按模板保留，后续继续补充。')

DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOC_OUT))
print(DOC_OUT)
print(REVIEW_MD)
print(COMMON_MD)
print(count_total)
