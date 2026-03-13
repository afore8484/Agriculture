from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r'E:/workspaces/Agriculture/prototype')
DOC_OUT = BASE / 'output' / 'doc' / 'spec_analysis_draft_v2_interfaces.docx'
PRINCIPLES_MD = BASE / 'docs' / 'interface_writing_principles.md'
CATALOG_MD = BASE / 'docs' / 'interface_catalog_outline.md'
CATALOG_JSON = BASE / 'data' / 'interfaces' / 'interface_catalog_seed.json'
PLACEHOLDER = '【待补充】本章节暂按模板保留，后续可结合详细需求、接口、数据设计和非功能指标继续完善。'

modules = [
    {
        'name': '农经一张图',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\1.农经一张图\农经一张图-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\1.农经一张图\农经一张图数据结构文档.md',
        'submodules': [
            ('首页总览', '查询、统计、地图联动、天气时间、导出'),
            ('财务一张图', '查询、分档统计、趋势分析、详情穿透、导出'),
            ('资产一张图', '查询、统计、排行、详情'),
            ('资源一张图', '查询、统计、排行、详情'),
            ('合同一张图', '查询、统计、应收应付明细、导出'),
            ('成员一张图', '查询、统计、结构分析、详情'),
            ('股权一张图', '查询、统计、分红明细、排行、导出'),
            ('预警监督', '规则配置、预警查询、处置、日志、导出'),
        ],
    },
    {
        'name': '村委财务事务管理系统',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\2.村委财务事务管理系统\村委财务事务管理系统-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\2.村委财务事务管理系统\村委财务事务管理系统.md',
        'submodules': [
            ('财务中心', '凭证、日记账、账户、模板、对账、查询、导入导出'),
            ('报表中心', '报表定义、报表查询、打印、导出、快照'),
            ('资产管理', '资产卡片、变更、折旧、盘点、处置'),
            ('合同管理', '合同登记、变更、收付款、验收、终止、续签'),
            ('在线审批', '流程定义、实例、任务、支付审批、历史查询'),
            ('银农直联', '账户同步、交易流水、余额、回单、接口日志'),
            ('基层党建', '组织、党员、课程、考试、学习记录'),
            ('期末处理与结账', '结转试算、月结、反结账、期间控制'),
        ],
    },
    {
        'name': '村集体经济收入台账',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\3.村集体经济收入台账\村集体经济收入台账-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\3.村集体经济收入台账\村集体经济收入台账数据结构文档.md',
        'submodules': [
            ('收入记录管理', '新增、修改、删除、查询、详情、导入导出'),
            ('信息收集统计', '统计、汇总、趋势、对比、导出'),
            ('区划汇总表', '区划筛选、分级汇总、钻取、导出'),
            ('收入增长预警', '规则配置、预警查询、同比环比分析、导出'),
            ('GDP目标管理', '目标设置、进度跟踪、对比分析'),
            ('发展思路与战略管理', '记录、查询、修改、版本留痕'),
            ('移动端填报', '简表录入、查询、草稿、提交'),
        ],
    },
    {
        'name': '村集体绩效评价管理系统',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\4.村集体绩效评价管理系统\村集体绩效评价管理系统-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\4.村集体绩效评价管理系统\村集体绩效评价管理系统数据结构文档.md',
        'submodules': [
            ('评价数据采集', '新增、导入、校验、提交'),
            ('数据审核', '审核、退回、复核、状态流转'),
            ('评价指标与模型', '指标配置、权重配置、规则管理'),
            ('评价结果与报告', '评分、报告生成、查看、导出'),
            ('统计分析', '区划汇总、排行、趋势、对比'),
            ('系统配置', '字典、年度参数、角色权限'),
        ],
    },
    {
        'name': '村集体智能审计系统',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\5.村集体智能审计系统\村集体智能审计系统-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\5.村集体智能审计系统\村集体智能审计系统数据结构文档.md',
        'submodules': [
            ('审计计划', '新增、变更、查询、下发'),
            ('审计方案', '编制、审批、任务分解、调整'),
            ('通知与承诺书', '生成、下发、回执、留痕'),
            ('资料调取', '调取单、提交、账套关联、附件管理'),
            ('疑点管理', '识别、查证、整改、状态跟踪'),
            ('底稿管理', '编制、审核、归档'),
            ('报告与征求意见', '草稿、反馈、复核、正式报告、归档'),
            ('决定书与统计报表', '决定书生成、统计、导出'),
            ('通用支撑', '附件、日志、流程实例、权限'),
        ],
    },
    {
        'name': '新型农业经营主体财务经营监管系统',
        'source_req': r'F:\待办\2026\三农金融\4.原型图\6.新型农业经营主体财务经营监管系统\新型农业经营主体财务经营监管系统-三端-需求梳理.md',
        'source_data': r'F:\待办\2026\三农金融\4.原型图\6.新型农业经营主体财务经营监管系统\新型农业经营主体财务经营监管系统数据结构文档.md',
        'submodules': [
            ('首页与三端驾驶舱', '卡片统计、图表查询、排行、预警'),
            ('经营主体管理', '主体档案、报表报送、审核'),
            ('成员账户管理', '成员、出资、转增、转让、继承、交易'),
            ('盈余管理', '补助量化、捐赠量化、公积金、盈余分配'),
            ('补贴与项目管理', '项目列表、申报、进度、资料补充'),
            ('示范申报与监测', '申报、审核、监测、状态流转'),
            ('商品/店铺/客户管理', '主数据维护、查询、停启用'),
            ('进销存管理', '采购、销售、退货、库存、调拨、统计'),
            ('金融服务与信贷管理', '产品查询、申请、审核、放款、还款、催收'),
            ('助农服务与行情', '服务目录、行情查询、同步'),
            ('通知公告与政策法规', '发布、查询、全文检索'),
        ],
    },
]

common_sections = [
    ('6.3.2.1 接口命名规范', '统一采用“业务域/子模块/动作”命名方式。查询类接口优先使用 GET，新增使用 POST，修改使用 PUT，删除使用 DELETE，审批与状态流转类可使用 POST/PUT。'),
    ('6.3.2.2 统一返回结构', '所有接口统一返回 code、message、data、requestId、timestamp。列表型接口的 data 内统一包含 pageNum、pageSize、total、list。'),
    ('6.3.2.3 鉴权与权限规则', '接口必须明确登录校验、角色权限、数据权限和导出权限。涉及区划、组织、主体的数据读取，必须结合当前用户可见范围进行过滤。'),
    ('6.3.2.4 查询与分页规则', '统一支持关键词、状态、区划、年度、月份、创建时间、更新时间等查询条件。分页参数统一为 pageNum、pageSize，排序参数统一为 orderBy、orderDirection。'),
    ('6.3.2.5 字段与字典规则', '入参和出参字段命名需与数据结构文档保持一致；状态字段、类型字段、枚举字段必须绑定数据字典或枚举说明。'),
    ('6.3.2.6 错误码与异常处理', '接口需定义参数错误、鉴权失败、权限不足、对象不存在、状态非法、外部接口失败、导出失败等错误码，并附场景说明。'),
    ('6.3.2.7 导入导出与附件规则', '导入接口必须提供模板说明和校验规则；导出接口必须说明导出范围、异步/同步方式和任务状态；附件接口必须明确文件类型、大小限制和存储策略。'),
    ('6.3.2.8 流程接口规则', '涉及审批、审核、复核、申报、整改等流转接口，必须说明前置状态、目标状态、操作角色、流程留痕和回退规则。'),
]

principles_md = '''# 接口编写原则说明

## 1. 编写目的
本说明用于统一海南数字“三农”金融服务平台接口文档的编写口径、接口命名方式、字段来源规则、状态流转规则和文档输出格式，确保后续 6 个模块的接口文档能够按同一标准编制、评审、开发和联调。

## 2. 编写依据
1. 原型图设计要求和各模块需求梳理文档。
2. 各模块数据结构文档和补充结构说明。
3. 需求梳理与数据结构映射类文档。
4. 当前项目已生成的需求规格说明书初稿。

## 3. 编写范围
1. 只编写接口文档，不重复编写需求分析正文。
2. 覆盖查询、新增、修改、删除、审核、审批、导入、导出、同步、日志、附件等接口。
3. 覆盖内部软件接口、跨模块接口和外部对接接口。

## 4. 排除项
1. 不在接口章节中重复写现状分析、用户分析、业务背景描述。
2. 不在接口章节中展开原型需求逐条说明。
3. 不把数据库完整表结构直接照搬为接口文档，接口只引用必要字段。

## 5. 编写原则
1. 以页面为入口：页面有展示必有查询接口，页面有操作必有写接口。
2. 以数据结构为约束：接口字段必须能映射到已有表结构或明确标注新增字段。
3. 以流程状态为边界：涉及审批、审核、整改、放款、申报等流程时，必须说明状态流转和角色限制。
4. 以复用优先：区划、字典、附件、导出、日志、分页、鉴权等能力优先抽为公共接口规范。
5. 以可联调为目标：每个接口必须说明调用方、提供方、请求方式、入参、出参、错误码和示例。

## 6. 文档结构要求
每个接口必须至少包含以下内容：
- 接口名称
- 接口用途
- 调用方/提供方
- 请求方式与 URL
- 鉴权要求
- 请求参数
- 响应参数
- 错误码
- 请求示例
- 响应示例
- 备注

## 7. 校验规则
1. 每个原型页面至少对应一组接口。
2. 每个核心业务表至少能找到对应消费接口。
3. 每个导出、上传、审批、审核动作必须有独立接口或明确复用关系。
4. 每个接口必须能回答“谁调用、为何调用、调用什么数据、数据从哪里来、调用后落在哪个页面或流程节点”五个问题。
'''

catalog = {
    'project': '海南数字“三农”金融服务平台',
    'doc_target': str(DOC_OUT),
    'modules': [
        {
            'name': module['name'],
            'source_req': module['source_req'],
            'source_data': module['source_data'],
            'submodules': [
                {'name': name, 'api_groups': groups}
                for name, groups in module['submodules']
            ],
        }
        for module in modules
    ],
}

catalog_lines = ['# 接口目录骨架', '', '## 1. 公共接口规范', '1. 登录与鉴权接口', '2. 区划与组织树接口', '3. 字典与枚举接口', '4. 附件上传下载接口', '5. 导入导出任务接口', '6. 日志与操作留痕接口', '']
for idx, module in enumerate(modules, 1):
    catalog_lines.append(f'## {idx + 1}. {module["name"]}')
    catalog_lines.append(f'- 原型来源：`{module["source_req"]}`')
    catalog_lines.append(f'- 数据来源：`{module["source_data"]}`')
    for sub_name, api_groups in module['submodules']:
        catalog_lines.append(f'- {sub_name}：{api_groups}')
    catalog_lines.append('')

PRINCIPLES_MD.parent.mkdir(parents=True, exist_ok=True)
CATALOG_MD.parent.mkdir(parents=True, exist_ok=True)
CATALOG_JSON.parent.mkdir(parents=True, exist_ok=True)
PRINCIPLES_MD.write_text(principles_md, encoding='utf-8')
CATALOG_MD.write_text('\n'.join(catalog_lines), encoding='utf-8')
CATALOG_JSON.write_text(json.dumps(catalog, ensure_ascii=False, indent=2), encoding='utf-8')


def set_cn_font(run, size=12, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def add_run(para, text, size=12, bold=False, center=False):
    run = para.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run


def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_cn_font(run, 10)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    size_map = {1: 16, 2: 14, 3: 12, 4: 12}
    add_run(p, text, size=size_map.get(level, 12), bold=True)


def paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, text, size=12)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.line_spacing = 1.5
        add_run(p, item, size=12)


def module_table(doc, module):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '子模块'
    hdr[1].text = '接口分组'
    hdr[2].text = '接口编写来源'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, 10, True)
    for sub_name, api_groups in module['submodules']:
        row = table.add_row().cells
        row[0].text = sub_name
        row[1].text = api_groups
        row[2].text = '原型需求梳理 + 数据结构文档'
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, 10, False)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
add_page_number(section)

add_run(doc.add_paragraph(), '海南数字“三农”金融服务平台', size=20, bold=True, center=True)
add_run(doc.add_paragraph(), '需求规格说明书（第二版）', size=18, bold=True, center=True)
add_run(doc.add_paragraph(), '适用范围：现状分析、用户分析、业务需求分析、接口编写原则与接口目录骨架', size=12, center=True)
add_run(doc.add_paragraph(), '版本：V2.0', size=12, center=True)
doc.add_paragraph('')
paragraph(doc, '说明：本版在第一版基础上新增“接口编写原则说明”“公共接口规范”和“6个模块接口目录骨架”。按你的要求，接口章节不回填需求分析正文，只保留接口编写依据、接口范围、接口规范和接口目录。')

heading(doc, '1. 引言', 1)
heading(doc, '1.1 编写目的', 2)
paragraph(doc, '本文档用于明确海南数字“三农”金融服务平台项目在现状分析、用户分析、业务需求分析和接口文档编制阶段的统一口径，为后续原型、详细接口、数据库、联调和测试工作提供依据。')
heading(doc, '1.2 项目背景', 2)
paragraph(doc, '在数字化转型持续深化的背景下，信息技术正在全面重塑农业农村领域的治理方式、生产方式和服务方式。随着乡村振兴战略深入实施，农业农村领域对数据驱动、精准治理和金融赋能的需求持续增强，传统以人工采集、分散管理、线下办理为主的业务模式已难以满足现代农业农村高质量发展的要求。')
paragraph(doc, '本项目以海南数字“三农”金融服务平台建设为抓手，围绕农业农村金融服务所需数据资源和业务场景，推进数据汇聚、治理、共享、分析、应用和监管一体化建设，形成“政府监管、金融服务、主体经营、村级治理”协同联动的数字化支撑体系。')
heading(doc, '1.3 术语与定义', 2)
paragraph(doc, '三农：指农业、农村和农民。')
paragraph(doc, '农经一张图：指对农业农村经济相关数据进行统一汇聚、统计分析、可视化展示和风险预警的综合监管应用。')
heading(doc, '1.4 参考资料', 2)
numbered(doc, [
    '《海南数字“三农”金融服务平台需求分析》相关材料。',
    'F盘原型图目录下 6 个模块的需求梳理文档与数据结构文档。',
    '当前项目已生成的需求规格说明书第一版。',
])

heading(doc, '2. 项目总体描述', 1)
heading(doc, '2.1 项目目标', 2)
numbered(doc, [
    '充分激活已有农业农村大数据资源，提升数据资产化和业务应用能力。',
    '构建面向政府部门、金融机构、村集体组织和新型农业经营主体的综合服务平台。',
    '提升金融支农服务的精准性、可获得性和效率，缓解融资难、融资贵、融资慢问题。',
    '强化农业农村领域数据治理、动态监管和风险预警能力。',
])
heading(doc, '2.2 用户特征', 2)
paragraph(doc, '平台面向多主体协同应用，覆盖政府农业管理部门、金融机构、村集体组织、新型农业经营主体及农业服务提供商等用户群体。')
heading(doc, '2.3 运行环境', 2)
paragraph(doc, PLACEHOLDER)
heading(doc, '2.4 约束与假设', 2)
paragraph(doc, PLACEHOLDER)

heading(doc, '3. 功能需求', 1)
heading(doc, '3.1 系统总体功能结构', 2)
paragraph(doc, '本项目围绕“数字三农金融数据底座 + 业务应用系统”进行建设，形成“数据汇聚、数据治理、监管分析、金融服务、业务协同”一体化功能体系。总体包括数字“三农”金融数据总平台以及农经一张图、村委财务事务管理系统、村集体经济收入台账、村集体绩效评价管理系统、村集体智能审计系统、新型农业经营主体财务经营监管系统六大业务模块。')
heading(doc, '3.2 功能模块划分', 2)
for module in modules:
    paragraph(doc, module['name'])
heading(doc, '3.3 详细功能说明', 2)
paragraph(doc, '本章节在第一版中已完成业务需求层说明，本版不再重复展开需求条目，接口相关内容统一写入第6章。')

heading(doc, '4. 数据需求', 1)
paragraph(doc, PLACEHOLDER)
heading(doc, '5. 非功能需求', 1)
paragraph(doc, PLACEHOLDER)

heading(doc, '6. 外部接口需求', 1)
paragraph(doc, '本章只描述接口编写依据、接口范围、接口规范和接口目录骨架，不重复写现状、用户和业务需求分析内容。后续详细接口正文将在本章基础上继续补齐。')
heading(doc, '6.1 用户界面（UI）要求', 2)
numbered(doc, [
    '本章节仅描述页面与接口的调用关系，不展开原型需求逐条说明。',
    '每个页面的展示区、筛选区、详情区、操作区、导出区都必须映射到对应接口。',
    '图表、卡片、列表、树、地图、弹窗等交互元素应在接口文档中明确数据来源和刷新时机。',
    'UI 章节中的接口说明以“页面 -> 接口分组 -> 关键字段”方式组织。',
])
heading(doc, '6.2 硬件接口', 2)
paragraph(doc, '当前项目以 Web 系统和服务接口为主，本阶段未识别到专用硬件控制接口。若后续接入读卡器、打印终端、扫描设备或物联网终端，应在本章节新增硬件接入协议、驱动依赖和设备异常处理说明。')
heading(doc, '6.3 软件接口', 2)
heading(doc, '6.3.1 接口编写原则说明', 3)
numbered(doc, [
    '接口编写以原型图设计要求为场景依据，以数据结构文档为字段依据，以需求与数据结构映射说明为归属依据，以需求规格说明书为交付格式依据。',
    '接口章节只写接口，不重复写需求分析正文。',
    '页面有展示必有查询接口，页面有操作必有写接口，页面有状态流转必有流程接口。',
    '接口字段必须能映射到数据结构文档中的现有表或明确标注为后续补充字段。',
    '每个接口必须明确调用方、提供方、请求方式、入参、出参、错误码和示例。',
    '跨模块公共能力优先抽成公共接口规范，不在每个模块内重复定义。',
])
heading(doc, '6.3.2 公共接口规范', 3)
for title, body in common_sections:
    heading(doc, title, 4)
    paragraph(doc, body)

start_idx = 3
for offset, module in enumerate(modules, start=start_idx):
    heading(doc, f'6.3.{offset} {module["name"]}接口目录', 3)
    paragraph(doc, f'原型来源：{module["source_req"]}')
    paragraph(doc, f'数据来源：{module["source_data"]}')
    paragraph(doc, '说明：本节当前输出接口目录骨架，后续按“接口名称、用途、URL、请求方式、入参、出参、错误码、示例”逐条展开。')
    module_table(doc, module)

heading(doc, '7. 系统部署与运行需求', 1)
paragraph(doc, PLACEHOLDER)
heading(doc, '7.1 部署架构', 2)
paragraph(doc, PLACEHOLDER)
heading(doc, '7.2 灾备与恢复策略', 2)
paragraph(doc, PLACEHOLDER)
heading(doc, '7.3 系统监控与运维支持', 2)
paragraph(doc, PLACEHOLDER)

DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOC_OUT))
print(DOC_OUT)
print(PRINCIPLES_MD)
print(CATALOG_MD)
print(CATALOG_JSON)
