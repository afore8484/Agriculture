from copy import deepcopy
from pathlib import Path
import shutil
import openpyxl
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\workspaces\Agriculture\prototype")
DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.docx"
BACKUP_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.before_33_reorder.docx"
XLSX_PATH = BASE / "tmp" / "spreadsheets" / "contract_compare_1.xlsx"


def insert_paragraph_after(anchor, text=""):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        para.add_run(text)
    return para


def clone_table_after(doc, anchor, data):
    sample = None
    for t in doc.tables:
        head = [c.text for c in t.rows[0].cells]
        if head == ['??', '????', '????', '????']:
            sample = t
            break
    if sample is None:
        raise RuntimeError('function table sample not found')
    new_tbl = deepcopy(sample._tbl)
    for tr in list(new_tbl.tr_lst):
        new_tbl.remove(tr)
    row_template = deepcopy(sample.rows[0]._tr)
    for _ in data:
        new_tbl.append(deepcopy(row_template))
    anchor._p.addnext(new_tbl)
    table = Table(new_tbl, anchor._parent)
    table.style = sample.style
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row[:4]):
            table.cell(r_idx, c_idx).text = str(value)
    return table


def delete_between(p_start, p_end):
    cur = p_start._p.getnext()
    while cur is not None and cur != p_end._p:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def find_after(paragraphs, prefix, start_idx=0):
    for i in range(start_idx, len(paragraphs)):
        if paragraphs[i].text.strip().startswith(prefix):
            return i
    raise ValueError(prefix)


def load_modules():
    wb = openpyxl.load_workbook(XLSX_PATH, data_only=True)
    ws = wb['Sheet1']
    modules = []
    mod_index = {}
    cur_module = None
    cur_sub = None
    for r in range(3, ws.max_row + 1):
        seq = ws.cell(r, 1).value
        m1 = ws.cell(r, 2).value
        m2 = ws.cell(r, 3).value
        c4 = ws.cell(r, 4).value
        c5 = ws.cell(r, 5).value
        c6 = ws.cell(r, 6).value
        c7 = ws.cell(r, 7).value
        c8 = ws.cell(r, 8).value
        c10 = ws.cell(r, 10).value
        if m1:
            cur_module = str(m1).strip()
        if m2:
            cur_sub = str(m2).strip()
        func = c4 or c8 or c7
        detail = c7 or c6 or c5 or c10 or ''
        if not cur_module or not cur_sub or not func:
            continue
        if cur_module not in mod_index:
            mod_index[cur_module] = len(modules)
            modules.append({'name': cur_module, 'subs': [], 'sub_map': {}})
        module = modules[mod_index[cur_module]]
        if cur_sub not in module['sub_map']:
            module['sub_map'][cur_sub] = len(module['subs'])
            module['subs'].append({'name': cur_sub, 'items': []})
        sub = module['subs'][module['sub_map'][cur_sub]]
        sub['items'].append({
            'seq': '' if seq is None else str(seq),
            'func': str(func).strip(),
            'detail': str(detail).strip() or str(func).strip(),
        })
    return modules


def summary_text(module):
    names = '?'.join(sub['name'] for sub in module['subs'])
    cnt = sum(len(sub['items']) for sub in module['subs'])
    return f'?????{name if False else names}??????'


def function_table(module):
    rows = [['??', '????', '????', '????']]
    for sub in module['subs']:
        for item in sub['items']:
            rows.append([item['seq'], sub['name'], item['func'], item['detail']])
    return rows


def main():
    modules = load_modules()
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)
    doc = Document(DOCX_PATH)
    paragraphs = doc.paragraphs
    i33 = find_after(paragraphs, '3.3 ', 80)
    i4 = find_after(paragraphs, '4. ', i33 + 1)
    p33 = paragraphs[i33]
    p4 = paragraphs[i4]
    delete_between(p33, p4)

    anchor = insert_paragraph_after(p33, '??????????????????????????????????')
    for idx, module in enumerate(modules, start=1):
        anchor = insert_paragraph_after(anchor, f'3.3.{idx} {module["name"]}????')
        names = '?'.join(sub['name'] for sub in module['subs'])
        anchor = insert_paragraph_after(anchor, f'?????{names}??????')
        table = clone_table_after(doc, anchor, function_table(module))
        # move anchor to a paragraph after table so the next insertion lands after the table
        anchor = insert_paragraph_after(anchor, '')
    doc.save(DOCX_PATH)
    print('reordered')


if __name__ == '__main__':
    main()
