from copy import deepcopy
from pathlib import Path
import json
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\workspaces\Agriculture\prototype")
DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.docx"
BACKUP_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.before_section_repair.docx"
PAYLOAD_PATH = BASE / "tmp" / "spec_repair_payload.json"


def insert_paragraph_after(anchor, text=""):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        para.add_run(text)
    return para


def clone_table_after(doc, anchor, data):
    sample = doc.tables[0]
    new_tbl = deepcopy(sample._tbl)
    for tr in list(new_tbl.tr_lst):
        new_tbl.remove(tr)
    row_template = deepcopy(sample.rows[0]._tr)
    for _ in data:
        new_tbl.append(deepcopy(row_template))
    anchor._p.addnext(new_tbl)
    table = Table(new_tbl, anchor._parent)
    table.style = sample.style
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row[:4]):
            table.cell(r_idx, c_idx).text = str(value)
    return table


def delete_between(p_start, p_end):
    cur = p_start._p.getnext()
    while cur is not None and cur != p_end._p:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def find_after(paragraphs, prefix, start_idx):
    for i in range(start_idx, len(paragraphs)):
        if paragraphs[i].text.strip().startswith(prefix):
            return i
    raise ValueError(prefix)


def main():
    payload = json.loads(PAYLOAD_PATH.read_text(encoding="utf-8-sig"))
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    paragraphs = doc.paragraphs

    i_231 = find_after(paragraphs, "2.3.1 ", 50)
    i_232 = find_after(paragraphs, "2.3.2 ", i_231 + 1)
    i_24 = find_after(paragraphs, "2.4 ", i_232 + 1)
    i_421 = find_after(paragraphs, "4.2.1 ", 120)
    i_422 = find_after(paragraphs, "4.2.2 ", i_421 + 1)
    i_43 = find_after(paragraphs, "4.3 ", i_422 + 1)
    i_44 = find_after(paragraphs, "4.4 ", i_43 + 1)
    i_5 = find_after(paragraphs, "5. ", i_44 + 1)

    p_231 = paragraphs[i_231]
    p_232 = paragraphs[i_232]
    p_24 = paragraphs[i_24]
    p_421 = paragraphs[i_421]
    p_422 = paragraphs[i_422]
    p_43 = paragraphs[i_43]
    p_44 = paragraphs[i_44]
    p_5 = paragraphs[i_5]

    delete_between(p_231, p_232)
    delete_between(p_232, p_24)
    delete_between(p_421, p_422)
    delete_between(p_43, p_44)
    delete_between(p_44, p_5)

    p = insert_paragraph_after(p_231, payload["hardware_intro"])
    clone_table_after(doc, p, payload["hardware_rows"])

    p = insert_paragraph_after(p_232, payload["software_intro"])
    clone_table_after(doc, p, payload["software_rows"])

    p = insert_paragraph_after(p_421, payload["db_intro"])
    clone_table_after(doc, p, payload["db_rows"])

    p = insert_paragraph_after(p_43, payload["dict_intro_1"])
    p = insert_paragraph_after(p, payload["dict_intro_2"])
    clone_table_after(doc, p, payload["dict_rows"])

    p = insert_paragraph_after(p_44, payload["collect_intro"])
    p = insert_paragraph_after(p, "4.4.1 ????")
    p = insert_paragraph_after(p, payload["collect_scope_intro"])
    clone_table_after(doc, p, payload["collect_scope_rows"])
    p = insert_paragraph_after(p, "4.4.2 ????????")
    p = insert_paragraph_after(p, payload["collect_source_intro"])
    clone_table_after(doc, p, payload["collect_source_rows"])
    p = insert_paragraph_after(p, "4.4.3 ????")
    p = insert_paragraph_after(p, payload["collect_strategy_intro"])
    clone_table_after(doc, p, payload["collect_strategy_rows"])
    p = insert_paragraph_after(p, "4.4.4 ???????????")
    p = insert_paragraph_after(p, payload["collect_quality_intro"])
    clone_table_after(doc, p, payload["collect_quality_rows"])
    insert_paragraph_after(p, payload["collect_note"])

    doc.save(DOCX_PATH)
    print("repaired")


if __name__ == "__main__":
    main()
