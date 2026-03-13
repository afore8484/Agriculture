from pathlib import Path
import openpyxl
from docx import Document

BASE = Path(r"E:\workspaces\Agriculture\prototype")
DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.docx"
XLSX_PATH = BASE / "tmp" / "spreadsheets" / "contract_compare_1.xlsx"

INTRO_32 = "???????????-??-???????????? 1.xlsx????????????????????????????????????????????????? xls ??????????"
INTRO_33 = "??? xls ??????????????????????? xls ????????????????????????????????"
SUMMARY_NOTE = "? xls ????????"
TABLE_SUMMARY_HEAD = ["????", "?????", "?????", "????"]
TABLE_FUNC_HEAD = ["??", "????", "????", "????"]


def load_modules():
    wb = openpyxl.load_workbook(XLSX_PATH, data_only=True)
    ws = wb['Sheet1']
    modules = []
    mod_index = {}
    cur_module = None
    cur_sub = None
    for r in range(3, ws.max_row + 1):
        m1 = ws.cell(r, 2).value
        m2 = ws.cell(r, 3).value
        c4 = ws.cell(r, 4).value
        c8 = ws.cell(r, 8).value
        c7 = ws.cell(r, 7).value
        if m1:
            cur_module = str(m1).strip()
        if m2:
            cur_sub = str(m2).strip()
        func = c4 or c8 or c7
        if not cur_module or not cur_sub or not func:
            continue
        if cur_module not in mod_index:
            mod_index[cur_module] = len(modules)
            modules.append({'name': cur_module, 'subs': [], 'sub_map': {}})
        module = modules[mod_index[cur_module]]
        if cur_sub not in module['sub_map']:
            module['sub_map'][cur_sub] = len(module['subs'])
            module['subs'].append({'name': cur_sub, 'count': 0})
        module['subs'][module['sub_map'][cur_sub]]['count'] += 1
    return modules


def module_summary_text(module):
    sub_names = '?'.join(sub['name'] for sub in module['subs'])
    item_count = sum(sub['count'] for sub in module['subs'])
    return f"???? xls ????????{len(module['subs'])}??????{item_count}??????????{sub_names}?"


def main():
    modules = load_modules()
    doc = Document(DOCX_PATH)

    # Fix paragraphs
    p32_idx = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip().startswith('3.2 '))
    p33_idx = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip().startswith('3.3 '))
    doc.paragraphs[p32_idx + 1].text = INTRO_32
    for i, module in enumerate(modules, start=1):
        idx = next(j for j,p in enumerate(doc.paragraphs) if p.text.strip().startswith(f'3.2.{i} '))
        doc.paragraphs[idx].text = f'3.2.{i} {module["name"]}'
        doc.paragraphs[idx + 1].text = module_summary_text(module)
    doc.paragraphs[p33_idx + 1].text = INTRO_33
    for i, module in enumerate(modules, start=1):
        idx = next(j for j,p in enumerate(doc.paragraphs) if p.text.strip().startswith(f'3.3.{i} '))
        doc.paragraphs[idx].text = f'3.3.{i} {module["name"]}????'
        doc.paragraphs[idx + 1].text = module_summary_text(module)

    # Fix tables
    # Summary table: detect by first data row module name
    for t in doc.tables:
        if len(t.rows) >= 2 and t.rows[1].cells[0].text == modules[0]['name']:
            for c, value in enumerate(TABLE_SUMMARY_HEAD):
                t.rows[0].cells[c].text = value
            for r, module in enumerate(modules, start=1):
                t.rows[r].cells[3].text = SUMMARY_NOTE
        if len(t.rows) >= 2 and t.rows[1].cells[0].text.isdigit():
            for c, value in enumerate(TABLE_FUNC_HEAD):
                t.rows[0].cells[c].text = value

    doc.save(DOCX_PATH)
    print('fixed')


if __name__ == '__main__':
    main()
