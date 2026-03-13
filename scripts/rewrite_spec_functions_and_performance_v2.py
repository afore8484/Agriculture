from copy import deepcopy
from pathlib import Path
import json
import shutil
import openpyxl
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\workspaces\Agriculture\prototype")
DOCX_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.docx"
BACKUP_PATH = BASE / "output" / "doc" / "spec_analysis_draft_v1.before_33_51_rewrite.docx"
XLSX_PATH = BASE / "tmp" / "spreadsheets" / "contract_compare_1.xlsx"
PAYLOAD_PATH = BASE / "tmp" / "rewrite_spec_payload.json"


def insert_paragraph_after(anchor, text=""):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        para.add_run(text)
    return para


def clone_table_after(doc, anchor, data):
    sample = doc.tables[0]
    new_tbl = deepcopy(sample._tbl)
    for tr in list(new_tbl.tr_lst):
        new_tbl.remove(tr)
    row_template = deepcopy(sample.rows[0]._tr)
    for _ in data:
        new_tbl.append(deepcopy(row_template))
    anchor._p.addnext(new_tbl)
    table = Table(new_tbl, anchor._parent)
    table.style = sample.style
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row[:4]):
            table.cell(r_idx, c_idx).text = str(value)
    return table


def delete_between(p_start, p_end):
    cur = p_start._p.getnext()
    while cur is not None and cur != p_end._p:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def parse_workbook():
    wb = openpyxl.load_workbook(XLSX_PATH, data_only=True)
    ws = wb['Sheet1']
    cur_mod1 = None
    cur_mod2 = None
    mods = {}
    for r in range(3, ws.max_row + 1):
        m1 = ws.cell(r, 2).value
        m2 = ws.cell(r, 3).value
        contract = ws.cell(r, 4).value
        bid = ws.cell(r, 7).value
        if m1:
            cur_mod1 = str(m1).strip()
        if m2:
            cur_mod2 = str(m2).strip()
        func = contract or bid
        if not cur_mod1 or not cur_mod2 or not func:
            continue
        mods.setdefault(cur_mod1, {}).setdefault(cur_mod2, []).append(str(func).strip())
    return mods


def top_examples(items, n=3):
    seen = []
    for item in items:
        if item not in seen:
            seen.append(item)
        if len(seen) >= n:
            break
    return '?'.join(seen)


def module_table(mods, module_name, payload):
    rows = [payload['module_table_header']]
    for sub, items in sorted(mods[module_name].items(), key=lambda kv: (-len(kv[1]), kv[0])):
        rows.append([
            sub,
            str(len(items)),
            top_examples(items),
            sub + payload['module_table_desc_suffix']
        ])
    return rows


def supervision_table(mods, supervision_modules, terminal_names, payload):
    rows = [payload['supervision_table_header']]
    for module_name in supervision_modules:
        for sub, items in sorted(mods[module_name].items(), key=lambda kv: (-len(kv[1]), kv[0])):
            rows.append([terminal_names[module_name], sub, str(len(items)), top_examples(items)])
    return rows


def find_after(paragraphs, prefix, start_idx):
    for i in range(start_idx, len(paragraphs)):
        if paragraphs[i].text.strip().startswith(prefix):
            return i
    raise ValueError(prefix)


def main():
    payload = json.loads(PAYLOAD_PATH.read_text(encoding="utf-8-sig"))
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    mods = parse_workbook()
    totals = {k: sum(len(v) for v in subs.values()) for k, subs in mods.items()}
    module_242 = next(k for k, v in totals.items() if v == 242)
    module_220 = next(k for k, v in totals.items() if v == 220)
    module_107 = next(k for k, v in totals.items() if v == 107)
    module_110 = next(k for k, v in totals.items() if v == 110)
    module_141 = next(k for k, v in totals.items() if v == 141)
    module_121 = next(k for k, v in totals.items() if v == 121)
    module_27 = next(k for k, v in totals.items() if v == 27)
    module_30 = next(k for k, v in totals.items() if v == 30)
    supervision_modules = [module_121, module_27, module_30]
    terminal_names = {module_121: '???', module_27: '???', module_30: '???'}
    supervision_total = sum(totals[k] for k in supervision_modules)

    doc = Document(DOCX_PATH)
    paragraphs = doc.paragraphs
    i331 = find_after(paragraphs, '3.3.1 ', 90)
    i332 = find_after(paragraphs, '3.3.2 ', i331 + 1)
    i333 = find_after(paragraphs, '3.3.3 ', i332 + 1)
    i334 = find_after(paragraphs, '3.3.4 ', i333 + 1)
    i335 = find_after(paragraphs, '3.3.5 ', i334 + 1)
    i336 = find_after(paragraphs, '3.3.6 ', i335 + 1)
    i337 = find_after(paragraphs, '3.3.7 ', i336 + 1)
    i4 = find_after(paragraphs, '4. ', i337 + 1)
    i51 = find_after(paragraphs, '5.1 ', 160)
    i52 = find_after(paragraphs, '5.2 ', i51 + 1)

    p331 = paragraphs[i331]
    p332 = paragraphs[i332]
    p333 = paragraphs[i333]
    p334 = paragraphs[i334]
    p335 = paragraphs[i335]
    p336 = paragraphs[i336]
    p337 = paragraphs[i337]
    p4 = paragraphs[i4]
    p51 = paragraphs[i51]
    p52 = paragraphs[i52]

    for start, end in [(p331, p332), (p332, p333), (p333, p334), (p334, p335), (p335, p336), (p336, p337), (p337, p4), (p51, p52)]:
        delete_between(start, end)

    p = insert_paragraph_after(p331, payload['section_331_intro'])
    clone_table_after(doc, p, payload['platform_rows'])

    p = insert_paragraph_after(p332, payload['section_332_intro'])
    clone_table_after(doc, p, module_table(mods, module_242, payload))

    p = insert_paragraph_after(p333, payload['section_333_intro'])
    clone_table_after(doc, p, module_table(mods, module_220, payload))

    p = insert_paragraph_after(p334, payload['section_334_intro'])
    clone_table_after(doc, p, module_table(mods, module_107, payload))

    p = insert_paragraph_after(p335, payload['section_335_intro'])
    clone_table_after(doc, p, module_table(mods, module_110, payload))

    p = insert_paragraph_after(p336, payload['section_336_intro'])
    clone_table_after(doc, p, module_table(mods, module_141, payload))

    p = insert_paragraph_after(p337, payload['section_337_intro_prefix'] + str(supervision_total) + payload['section_337_intro_suffix'])
    clone_table_after(doc, p, supervision_table(mods, supervision_modules, terminal_names, payload))

    p = insert_paragraph_after(p51, payload['performance_intro'])
    p = insert_paragraph_after(p, payload['performance_heading_1'])
    clone_table_after(doc, p, payload['performance_rows'])
    p = insert_paragraph_after(p, payload['performance_heading_2'])
    clone_table_after(doc, p, payload['performance_condition_rows'])
    insert_paragraph_after(p, payload['performance_note'])

    doc.save(DOCX_PATH)
    print('updated')


if __name__ == '__main__':
    main()
